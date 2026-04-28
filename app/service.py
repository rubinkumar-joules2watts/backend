from __future__ import annotations

from copy import deepcopy
from datetime import datetime, timedelta, timezone
import json
import logging
import os
from pathlib import Path
import re
import shutil
from typing import Any
from uuid import uuid4

import httpx

from fastapi import HTTPException, UploadFile
from pymongo.database import Database

from .config import COLLECTIONS, UPLOADS_DIR
from .ai import insight_service


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


def ensure_table(table: str) -> None:
    if table not in COLLECTIONS:
        raise HTTPException(status_code=404, detail=f"Unknown table: {table}")


def to_plain_document(document: dict[str, Any] | None) -> dict[str, Any] | None:
    if document is None:
        return None

    result = dict(document)
    result.pop("_id", None)
    result["id"] = result.get("id") or str(document.get("_id"))
    return result


def normalize_on_insert(table: str, document: dict[str, Any]) -> dict[str, Any]:
    next_document = deepcopy(document)
    now = utc_now_iso()
    next_document["id"] = str(next_document.get("id") or uuid4())

    if table == "clients":
        next_document["created_at"] = next_document.get("created_at") or now

    if table in {"projects", "milestones", "team_members"}:
        next_document["created_at"] = next_document.get("created_at") or now
        next_document["updated_at"] = next_document.get("updated_at") or now

    if table in {"project_assignments", "project_updates", "project_documents"}:
        next_document["created_at"] = next_document.get("created_at") or now

    if table == "audit_log":
        next_document["created_at"] = next_document.get("created_at") or now
        next_document["changed_by"] = next_document.get("changed_by") or "system"

    if table == "team_members_engagement":
        next_document["created_at"] = next_document.get("created_at") or now
        next_document["updated_at"] = next_document.get("updated_at") or now

    # Initialize week data arrays for milestones (will be populated after insert)
    if table == "milestones":
        next_document["practice_weeks"] = next_document.get("practice_weeks") or []
        next_document["signoff_weeks"] = next_document.get("signoff_weeks") or []
        next_document["invoice_weeks"] = next_document.get("invoice_weeks") or []

    return next_document


def upsert_team_member_engagement(database: Database, payload: Any) -> dict[str, Any]:
    """
    Upsert into `team_members_engagement` by (team_member_id, project_id).
    - If a record exists for the pair, update it.
    - Otherwise create a new record.
    """
    if not isinstance(payload, dict):
        raise HTTPException(status_code=400, detail="Request body must be a JSON object")

    team_member_id = str(payload.get("team_member_id") or "").strip()
    project_id = str(payload.get("project_id") or "").strip()
    if not team_member_id or not project_id:
        raise HTTPException(status_code=400, detail="team_member_id and project_id are required")

    engagement_collection = database["team_members_engagement"]
    existing = engagement_collection.find_one({"team_member_id": team_member_id, "project_id": project_id})

    if existing:
        existing_id = str(existing.get("id") or "")
        next_changes = deepcopy(payload)
        next_changes.pop("_id", None)
        next_changes.pop("id", None)
        next_changes["team_member_id"] = team_member_id
        next_changes["project_id"] = project_id
        next_changes["updated_at"] = utc_now_iso()

        engagement_collection.update_one({"id": existing_id}, {"$set": next_changes})
        updated = engagement_collection.find_one({"id": existing_id})
        return to_plain_document(updated) or {}

    document = normalize_on_insert("team_members_engagement", payload)
    document["team_member_id"] = team_member_id
    document["project_id"] = project_id
    engagement_collection.insert_one(document)
    return to_plain_document(document) or {}


def build_filter(query_params: dict[str, str]) -> dict[str, str]:
    filter_doc: dict[str, str] = {}
    for key, raw_value in query_params.items():
        if key in {"orderBy", "ascending", "limit", "offset"}:
            continue
        if raw_value in {"", None}:
            continue
        filter_doc[key] = str(raw_value)
    return filter_doc


def parse_date(date_str: str | None) -> datetime | None:
    """Parse date string to datetime object. Handles multiple formats: YYYY-MM-DD, ISO format."""
    if not date_str:
        return None
    try:
        if isinstance(date_str, str):
            # Handle ISO format with Z
            if "T" in date_str:
                dt = datetime.fromisoformat(date_str.replace("Z", "+00:00"))
            # Handle simple date format YYYY-MM-DD
            elif "-" in date_str:
                dt = datetime.fromisoformat(date_str)
            else:
                return None

            # Ensure timezone-aware
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            return dt
        return date_str
    except (ValueError, AttributeError, TypeError):
        return None


def get_week_number_and_label(date: datetime, start_date: datetime) -> tuple[int, str]:
    """Get week number and label based on the actual date provided."""
    days_diff = (date.date() - start_date.date()).days
    week_number = days_diff // 7

    # Label the week by the date provided, not the week start
    # This ensures Feb 2 is labeled as "Feb 2, 2025" not "Jan 27, 2025"
    week_label = date.strftime("%b %d, %Y").lstrip("0").replace(" 0", " ")

    return week_number, week_label


def calculate_status(date: datetime, eta_date: datetime | None, actual_date: datetime | None) -> tuple[str, str]:
    """
    Calculate status and color based on date comparison.
    Returns (status, color).
    """
    # Make date timezone-aware if it's naive
    if date and date.tzinfo is None:
        date = date.replace(tzinfo=timezone.utc)

    if actual_date:
        # Make actual_date timezone-aware if it's naive
        if actual_date.tzinfo is None:
            actual_date = actual_date.replace(tzinfo=timezone.utc)

        if eta_date:
            # Make eta_date timezone-aware if it's naive
            if eta_date.tzinfo is None:
                eta_date = eta_date.replace(tzinfo=timezone.utc)

            if actual_date.date() <= eta_date.date():
                return "On Track", "green"
            else:
                days_diff = (actual_date.date() - eta_date.date()).days
                if days_diff <= 14:  # Within 2 weeks of ETA
                    return "At Risk", "orange"
                else:
                    return "At Risk", "orange"
        return "On Track", "green"
    else:
        now = datetime.now(timezone.utc)
        if date.date() <= now.date():
            return "Blocked", "red"
        days_until = (date.date() - now.date()).days
        if days_until <= 14:  # Within 2 weeks
            return "At Risk", "orange"
        else:
            return "On Track", "green"


def get_status_and_color(status_value: str | None, is_pending: bool = False) -> tuple[str, str]:
    """
    Determine status and color based on status value and pending flag.
    - Pending: Amber
    - Completed/Done: Blue
    - On Hold: Amber
    - On Track: Green
    - At Risk: Orange
    - Blocked: Red
    """
    if is_pending or status_value == "Pending":
        return "Pending", "amber"

    status_lower = (status_value or "").lower()

    if status_lower in {"completed", "done"}:
        return "Completed", "blue"
    elif status_lower == "on hold":
        return "On Hold", "amber"
    elif status_lower == "blocked":
        return "Blocked", "red"
    elif status_lower == "on track":
        return "On Track", "green"
    elif status_lower == "at risk":
        return "At Risk", "orange"

    return "On Track", "green"


def generate_practice_weeks(milestone: dict[str, Any], start_date: datetime, weeks_data: dict[int, dict[str, Any]]) -> list[dict[str, Any]]:
    """
    Generate practice week data based on actual_start, actual_end_eta, and status.
    Returns list of week markers with status and color for each week.

    Status Logic:
    - If status="Completed" AND actual_end_eta date has passed → Show as Completed (blue)
    - Else use the milestone status field as-is
    """
    practice_weeks = []

    actual_start = parse_date(milestone.get("actual_start"))
    actual_end = parse_date(milestone.get("actual_end_eta"))
    milestone_status = (milestone.get("status") or "").strip()

    if not actual_start or not actual_end:
        return practice_weeks

    # ============ DETERMINE FINAL WEEK STATUS AND COLOR ============
    # This is where we check the "Completed" status based on the milestone status field
    # The final week gets the current milestone status, regardless of the actual_end_eta date

    status_lower = milestone_status.lower() if milestone_status else ""

    # Map status to color and display status
    # - If status="Completed" → Show as blue/Completed
    # - If status="At Risk" → Show as orange/At Risk
    # - If status="Blocked" → Show as red/Blocked
    # - Otherwise → Show as green/On Track

    if status_lower == "completed":
        # ✅ MILESTONE IS MARKED AS COMPLETED
        final_color = "blue"
        final_status = "Completed"
    elif status_lower == "at risk":
        final_color = "orange"
        final_status = "At Risk"
    elif status_lower == "blocked":
        final_color = "red"
        final_status = "Blocked"
    else:  # "on track" or default
        final_color = "green"
        final_status = "On Track"

    # Find the Monday of the week containing actual_start
    week_start_of_start = actual_start - timedelta(days=actual_start.weekday())

    # If actual_start is on Fri/Sat/Sun (last 3 days of week), skip partial week
    if actual_start.weekday() >= 4:  # Friday=4, Saturday=5, Sunday=6
        week_start_of_start = week_start_of_start + timedelta(days=7)

    milestone_start_week_num = (week_start_of_start.date() - start_date.date()).days // 7

    # Find the Monday of the week containing actual_end
    week_start_of_end = actual_end - timedelta(days=actual_end.weekday())
    end_week_num = (week_start_of_end.date() - start_date.date()).days // 7

    # If milestone_start > end, include at least the end week
    if milestone_start_week_num > end_week_num:
        milestone_start_week_num = end_week_num

    # Include all weeks where the milestone is running
    for week_num_iter in range(milestone_start_week_num, end_week_num + 1):
        if week_num_iter in weeks_data:
            week_label = weeks_data[week_num_iter]["label"]

            is_final_week = week_num_iter == end_week_num
            is_first_week = week_num_iter == milestone_start_week_num

            if is_final_week:
                week_status = final_status
                week_color = final_color
            else:
                week_status = "On Track"
                week_color = "green"

            # Determine the date for this week
            if is_first_week and is_final_week:
                week_date = actual_end.isoformat()
            elif is_first_week:
                week_date = actual_start.isoformat()
            elif is_final_week:
                week_date = actual_end.isoformat()
            else:
                week_date = weeks_data[week_num_iter]["start"]

            practice_weeks.append({
                "week_number": week_num_iter,
                "week_label": week_label,
                "status": week_status,
                "color": week_color,
                "date": week_date
            })

    return practice_weeks


def generate_signoff_weeks(milestone: dict[str, Any], start_date: datetime, weeks_data: dict[int, dict[str, Any]]) -> list[dict[str, Any]]:
    """
    Generate signoff week data based on client_signoff_status.
    Returns list with single week marker showing signoff status.

    Logic:
    - If status="Done" → Use signedoff_date or actual_end_eta
    - If status="Pending" → Use actual_end_eta
    """
    signoff_weeks = []

    signoff_status_raw = (milestone.get("client_signoff_status") or "").strip().lower()
    actual_end_eta = parse_date(milestone.get("actual_end_eta"))

    # Determine which date to use based on status
    if signoff_status_raw == "done":
        signoff_date = parse_date(milestone.get("signedoff_date"))
        signoff_display_date = signoff_date or actual_end_eta
        signoff_status = "Done"
        signoff_color = "green"
    else:
        signoff_display_date = actual_end_eta
        signoff_status = "Pending"
        signoff_color = "orange"

    if signoff_display_date:
        week_num, _ = get_week_number_and_label(signoff_display_date, start_date)
        week_label = weeks_data[week_num]["label"] if week_num in weeks_data else ""

        signoff_weeks.append({
            "week_number": week_num,
            "week_label": week_label,
            "status": signoff_status,
            "color": signoff_color,
            "date": signoff_display_date.isoformat()
        })

    return signoff_weeks


def generate_invoice_weeks(milestone: dict[str, Any], start_date: datetime, weeks_data: dict[int, dict[str, Any]]) -> list[dict[str, Any]]:
    """
    Generate invoice week data based on invoice_status.
    Returns list with single week marker showing invoice status.

    Logic:
    - If status="Done" → Use invoice_raised_date or actual_end_eta
    - If status="Pending" → Use actual_end_eta
    """
    invoice_weeks = []

    invoice_status_raw = (milestone.get("invoice_status") or "").strip().lower()
    actual_end_eta = parse_date(milestone.get("actual_end_eta"))

    # Determine which date to use based on status
    if invoice_status_raw == "done":
        invoice_date = parse_date(milestone.get("invoice_raised_date"))
        invoice_display_date = invoice_date or actual_end_eta
        invoice_status = "Done"
        invoice_color = "green"
    else:
        invoice_display_date = actual_end_eta
        invoice_status = "Pending"
        invoice_color = "orange"

    if invoice_display_date:
        week_num, _ = get_week_number_and_label(invoice_display_date, start_date)
        week_label = weeks_data[week_num]["label"] if week_num in weeks_data else ""

        invoice_weeks.append({
            "week_number": week_num,
            "week_label": week_label,
            "status": invoice_status,
            "color": invoice_color,
            "date": invoice_display_date.isoformat()
        })

    return invoice_weeks


def generate_all_milestone_weeks(database: Database, milestone: dict[str, Any]) -> dict[str, Any]:
    """
    Generate and store all week data (practice, signoff, invoice) for a milestone.
    Called when a milestone is created or updated.
    """
    project_id = milestone.get("project_id")

    # Get all milestones to calculate date range
    all_milestones = list(database["milestones"].find({"project_id": str(project_id)}))

    all_dates = []
    for m in all_milestones:
        actual_start = parse_date(m.get("actual_start"))
        actual_end = parse_date(m.get("actual_end_eta"))
        signoff_date = parse_date(m.get("signedoff_date"))
        signoff_eta = parse_date(m.get("actual_end_eta"))
        invoice_date = parse_date(m.get("invoice_raised_date"))
        invoice_eta = parse_date(m.get("actual_end_eta"))

        for date in [actual_start, actual_end, signoff_date, signoff_eta, invoice_date, invoice_eta]:
            if date:
                all_dates.append(date)

    if not all_dates:
        return {
            "practice_weeks": [],
            "signoff_weeks": [],
            "invoice_weeks": []
        }

    start_date = min(all_dates)
    end_date = max(all_dates)

    # Extend range to include full weeks
    start_date = start_date - timedelta(days=start_date.weekday())
    end_date = end_date + timedelta(days=(6 - end_date.weekday()))

    # Generate all weeks
    weeks_data: dict[int, dict[str, Any]] = {}
    current = start_date
    week_num = 0
    while current <= end_date:
        week_end = current + timedelta(days=6)
        if current.month == week_end.month:
            week_label = f"{current.strftime('%b %d')}-{week_end.strftime('%d, %Y')}".lstrip("0").replace(" 0", " ")
        else:
            week_label = f"{current.strftime('%b %d')}-{week_end.strftime('%b %d, %Y')}".lstrip("0").replace(" 0", " ")
        weeks_data[week_num] = {"label": week_label, "start": current.isoformat()}
        current += timedelta(days=7)
        week_num += 1

    # Generate week data for each type
    practice_weeks = generate_practice_weeks(milestone, start_date, weeks_data)
    signoff_weeks = generate_signoff_weeks(milestone, start_date, weeks_data)
    invoice_weeks = generate_invoice_weeks(milestone, start_date, weeks_data)

    return {
        "practice_weeks": practice_weeks,
        "signoff_weeks": signoff_weeks,
        "invoice_weeks": invoice_weeks
    }


def regenerate_project_milestone_weeks(database: Database, project_id: str) -> None:
    """
    Regenerate week data for all milestones in a project.
    Called after creating or updating any milestone.
    """
    milestones = list(database["milestones"].find({"project_id": str(project_id)}))

    for milestone in milestones:
        week_data = generate_all_milestone_weeks(database, milestone)
        database["milestones"].update_one(
            {"id": milestone.get("id")},
            {"$set": week_data}
        )


def _count_weeks_in_month(year: int, month: int) -> int:
    """
    Count the number of Sun-Sat week rows displayed in a standard calendar view.
    Each week starts on Sunday and contains 7 consecutive days.
    
    For example:
    - February 2026 (starts on Sun): 4 week rows (1-7, 8-14, 15-21, 22-28)
    - March 2026 (starts on Sun): 5 week rows (1-7, 8-14, 15-21, 22-28, 29-31+)
    """
    first_day = datetime(year, month, 1)
    
    # Find the Sunday on or before the 1st of the month
    # In Python: weekday() returns Mon=0, Tue=1, ..., Sun=6
    # To find preceding Sunday, we calculate how many days back
    days_to_sunday = (first_day.weekday() + 1) % 7
    week_start = first_day - timedelta(days=days_to_sunday)
    
    # Get the last day of the month
    if month == 12:
        last_day = datetime(year + 1, 1, 1) - timedelta(days=1)
    else:
        last_day = datetime(year, month + 1, 1) - timedelta(days=1)
    
    # Count how many complete weeks (Sun-Sat) fit between week_start and last_day
    weeks_count = 0
    current_week_start = week_start
    while current_week_start <= last_day:
        weeks_count += 1
        current_week_start += timedelta(days=7)
    
    return weeks_count


_UUID_LIKE_RE = re.compile(
    r"^[0-9a-fA-F]{8}-"
    r"[0-9a-fA-F]{4}-"
    r"[0-9a-fA-F]{4}-"
    r"[0-9a-fA-F]{4}-"
    r"[0-9a-fA-F]{12}$"
)


def _resolve_project_display_name(project_doc: dict[str, Any] | None, fallback_project_id: str | None = None) -> str:
    """
    Resolve the best human-readable project name from mixed project schemas.
    Prefers explicit project-name fields over generic `name`.
    """
    if not project_doc:
        return fallback_project_id or "Unnamed Project"

    candidate_keys = (
        "project_name",
        "project_title",
        "projectTitle",
        "title",
        "name",
    )
    for key in candidate_keys:
        value = project_doc.get(key)
        if not value:
            continue
        text = str(value).strip()
        if not text:
            continue
        # Avoid showing UUIDs as display names.
        if _UUID_LIKE_RE.match(text):
            continue
        return text

    return fallback_project_id or "Unnamed Project"


def generate_calendar_months(start_date: datetime) -> list[dict[str, Any]]:
    """
    Generate month entries from start_date's month through December 2026.
    Each entry includes the count of Mon-Sun weeks whose Monday falls in that month.
    """
    if start_date.tzinfo is None:
        start_date = start_date.replace(tzinfo=timezone.utc)
    current = start_date.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    end = datetime(2026, 12, 1, tzinfo=timezone.utc)
    months = []
    while current <= end:
        months.append({
            "month": current.month,
            "year": current.year,
            "month_name": current.strftime("%B"),
            "month_year": current.strftime("%b %Y"),
            "weeks_count": _count_weeks_in_month(current.year, current.month),
        })
        current = (current.replace(month=current.month + 1) if current.month < 12
                   else current.replace(year=current.year + 1, month=1))
    return months


def get_milestone_health(database: Database, project_id: str) -> dict[str, Any]:
    """
    Get milestone health tracker data for a project.
    Uses pre-calculated week data stored in milestones collection.

    Filtering Logic:
    - Practice weeks: Show if actual_end_eta exists
    - Signoff: Show if actual_end_eta exists (calculate status based on client_signoff_status)
    - Invoice: Show if actual_end_eta exists (calculate status based on invoice_status)
    """
    # Fetch project
    project = database["projects"].find_one({"id": str(project_id)})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # Fetch all milestones for this project
    milestones = list(database["milestones"].find({"project_id": str(project_id)}))
    if not milestones:
        raise HTTPException(status_code=404, detail="No milestones found for this project")

    # Collect planned_start dates to build calendar month range
    # Falls back to actual_start if planned_start is not available
    planned_starts: list[datetime] = []
    for m in milestones:
        ps = parse_date(m.get("planned_start") or m.get("plannedStart") or m.get("actual_start"))
        if ps:
            planned_starts.append(ps)

    calendar_start_date = min(planned_starts) if planned_starts else datetime.now(timezone.utc)
    calendar_months = generate_calendar_months(calendar_start_date)

    # Collect all week data from milestones
    practice_data = []
    signoff_data = []
    invoice_data = []
    all_weeks_collected: dict[int, dict[str, Any]] = {}

    for milestone in milestones:
        milestone_code = milestone.get("milestone_code", "")
        description = milestone.get("description", "")
        milestone_id = milestone.get("id")
        actual_end_eta = milestone.get("actual_end_eta")

        # Get practice weeks (from stored data)
        # Show practice weeks only if actual_end_eta exists
        practice_weeks = milestone.get("practice_weeks", [])
        if practice_weeks and actual_end_eta:
            # Determine overall color based on final week status
            final_week = practice_weeks[-1]
            practice_color = final_week.get("color", "gray")

            # Collect weeks for all_weeks map
            for week in practice_weeks:
                week_num = week.get("week_number", 0)
                if week_num not in all_weeks_collected:
                    all_weeks_collected[week_num] = {
                        "label": week.get("week_label", ""),
                        "start": week.get("date", "")
                    }

            actual_start = parse_date(milestone.get("actual_start"))
            actual_end = parse_date(milestone.get("actual_end_eta"))

            practice_data.append({
                "id": milestone_id,
                "milestone_code": milestone_code,
                "description": description,
                "milestone_type": "practice",
                "start_date": actual_start.isoformat() if actual_start else None,
                "end_date": actual_end.isoformat() if actual_end else None,
                "weeks": practice_weeks,
                "completion_pct": milestone.get("completion_pct", 0),
                "status": milestone.get("status", ""),
                "color": practice_color,
                "days_variance": milestone.get("days_variance", 0)
            })

        # Get signoff data
        # Show signoff if actual_end_eta exists, regardless of whether signoff_weeks exists
        if actual_end_eta:
            signoff_weeks = milestone.get("signoff_weeks", [])
            
            # Collect weeks for all_weeks map if they exist
            for week in signoff_weeks:
                week_num = week.get("week_number", 0)
                if week_num not in all_weeks_collected:
                    all_weeks_collected[week_num] = {
                        "label": week.get("week_label", ""),
                        "start": week.get("date", "")
                    }

            signoff_display_date = parse_date(milestone.get("signedoff_date")) or parse_date(milestone.get("actual_end_eta"))
            
            # Calculate signoff status based on client_signoff_status
            client_signoff_status = milestone.get("client_signoff_status", "")
            if client_signoff_status == "Done":
                calculated_signoff_status = "Done"
            elif client_signoff_status == "Partial":
                calculated_signoff_status = "Partial"
            else:
                calculated_signoff_status = "Pending"

            signoff_data.append({
                "id": milestone_id,
                "milestone_code": milestone_code,
                "description": description,
                "milestone_type": "signoff",
                "date": signoff_display_date.isoformat() if signoff_display_date else None,
                "signoff_status": client_signoff_status,
                "status": calculated_signoff_status
            })

        # Get invoice data
        # Show invoice if actual_end_eta exists, regardless of whether invoice_weeks exists
        if actual_end_eta:
            invoice_weeks = milestone.get("invoice_weeks", [])
            
            # Collect weeks for all_weeks map if they exist
            for week in invoice_weeks:
                week_num = week.get("week_number", 0)
                if week_num not in all_weeks_collected:
                    all_weeks_collected[week_num] = {
                        "label": week.get("week_label", ""),
                        "start": week.get("date", "")
                    }

            invoice_display_date = parse_date(milestone.get("invoice_raised_date")) or parse_date(milestone.get("actual_end_eta"))
            
            # Calculate invoice status based on invoice_status
            invoice_status = milestone.get("invoice_status", "")
            if invoice_status == "Done":
                calculated_invoice_status = "Done"
            elif invoice_status == "Partial":
                calculated_invoice_status = "Partial"
            else:
                calculated_invoice_status = "Pending"

            invoice_data.append({
                "id": milestone_id,
                "milestone_code": milestone_code,
                "description": description,
                "milestone_type": "invoice",
                "date": invoice_display_date.isoformat() if invoice_display_date else None,
                "invoice_status": invoice_status,
                "status": calculated_invoice_status
            })

    # If no weeks were collected, return empty
    if not all_weeks_collected:
        return {
            "project_id": project_id,
            "project_name": project.get("name", ""),
            "practice": practice_data,
            "signoff": signoff_data,
            "invoice": invoice_data,
            "weeks_range": {
                "start_week": "",
                "end_week": "",
                "total_weeks": 0
            },
            "all_weeks": {},
            "calendar_start": calendar_start_date.strftime("%b %Y"),
            "calendar_months": calendar_months,
        }

    # Sort weeks by number
    sorted_week_nums = sorted(all_weeks_collected.keys())
    weeks_range = {
        "start_week": all_weeks_collected[sorted_week_nums[0]]["label"] if sorted_week_nums else "",
        "end_week": all_weeks_collected[sorted_week_nums[-1]]["label"] if sorted_week_nums else "",
        "total_weeks": len(all_weeks_collected)
    }

    return {
        "project_id": project_id,
        "project_name": project.get("name", ""),
        "practice": practice_data,
        "signoff": signoff_data,
        "invoice": invoice_data,
        "weeks_range": weeks_range,
        "all_weeks": all_weeks_collected,
        "calendar_start": calendar_start_date.strftime("%b %Y"),
        "calendar_months": calendar_months,
    }


def get_dashboard_counters(database: Database) -> dict[str, int]:
    """Return dashboard counts for project status counters."""
    projects = database["projects"]

    active_filter = {
        "$and": [
            {"status": {"$exists": True}},
            {"status": {"$not": {"$regex": "^(Blocked|Completed)$", "$options": "i"}}},
        ]
    }
    on_track_filter = {"status": {"$regex": "^On Track$", "$options": "i"}}
    at_risk_filter = {"status": {"$regex": "^At Risk$", "$options": "i"}}
    blocked_filter = {"status": {"$regex": "^Blocked$", "$options": "i"}}
    completed_filter = {"status": {"$regex": "^Completed$", "$options": "i"}}

    return {
        "total_projects": projects.count_documents({}),
        "active_projects": projects.count_documents(active_filter),
        "on_track_projects": projects.count_documents(on_track_filter),
        "at_risk_projects": projects.count_documents(at_risk_filter),
        "blocked_projects": projects.count_documents(blocked_filter),
        "completed_projects": projects.count_documents(completed_filter),
    }


def get_team_members_with_engagement(database: Database) -> list[dict[str, Any]]:
    """
    Get all team members with their engagement metrics aggregated from team_members_engagement collection.
    
    Returns:
        List of team members with engagement_pct and other engagement data
    """
    team_members_collection = database["team_members"]
    engagement_collection = database["team_members_engagement"]
    
    # Get all team members
    team_members = [to_plain_document(doc) for doc in team_members_collection.find({})]

    # Build project lookup once so each engagement can include project name
    projects_collection = database["projects"]
    project_name_by_id: dict[str, str] = {}
    for project in projects_collection.find({}, {"id": 1, "name": 1, "project_name": 1, "project_title": 1, "projectTitle": 1, "title": 1}):
        project_id = str(project.get("id") or "")
        if not project_id:
            continue
        project_name_by_id[project_id] = _resolve_project_display_name(project, None)
    
    # Enhance each team member with engagement data
    result = []
    for member in team_members:
        member_id = member.get("id")
        
        # Get all engagement records for this member
        engagements = list(engagement_collection.find({"team_member_id": member_id}))
        
        if engagements:
            enriched_engagements: list[dict[str, Any]] = []
            for eng in engagements:
                plain_eng = to_plain_document(eng) or {}
                project_id = str(plain_eng.get("project_id") or "")
                if project_id:
                    plain_eng["project_name"] = project_name_by_id.get(project_id, "Unnamed Project")
                else:
                    plain_eng["project_name"] = "Unnamed Project"

                # Skip unresolved projects from response payload.
                if plain_eng.get("project_name") == "Unnamed Project":
                    continue
                enriched_engagements.append(plain_eng)

            # engagement_pct should represent total engagement load per member.
            total_engagement_pct = 0.0
            for eng in enriched_engagements:
                level = eng.get("engagement_percentage")
                if level is None:
                    level = eng.get("engagement_level")
                try:
                    total_engagement_pct += float(level)
                except (ValueError, TypeError):
                    pass

            # Calculate totals only for valid/visible engagements.
            total_hours = sum(eng.get("engagement_hours", 0) for eng in enriched_engagements)
            total_tasks_completed = sum(eng.get("task_completed", 0) for eng in enriched_engagements)
            total_tasks_pending = sum(eng.get("task_pending", 0) for eng in enriched_engagements)
            project_count = len(enriched_engagements)

            # Add engagement data to member
            member["engagement_pct"] = round(total_engagement_pct, 2)
            member["total_engagement_hours"] = total_hours
            member["total_tasks_completed"] = total_tasks_completed
            member["total_tasks_pending"] = total_tasks_pending
            member["projects_assigned"] = project_count
            member["engagements"] = enriched_engagements
        else:
            # No engagements found
            member["engagement_pct"] = 0
            member["total_engagement_hours"] = 0
            member["total_tasks_completed"] = 0
            member["total_tasks_pending"] = 0
            member["projects_assigned"] = 0
            member["engagements"] = []
        
        result.append(member)

    # Deduplicate by member name so repeated records are merged into one card/user.
    deduped_by_name: dict[str, dict[str, Any]] = {}
    deduped_order: list[str] = []
    for member in result:
        name_raw = str(member.get("name") or "").strip()
        dedupe_key = " ".join(name_raw.lower().split()) if name_raw else f"id::{member.get('id')}"

        if dedupe_key not in deduped_by_name:
            base = deepcopy(member)
            base["engagements"] = list(member.get("engagements") or [])
            deduped_by_name[dedupe_key] = base
            deduped_order.append(dedupe_key)
            continue

        existing = deduped_by_name[dedupe_key]

        # Fill missing basic profile fields from duplicate rows.
        for key, value in member.items():
            if key == "engagements":
                continue
            if existing.get(key) in (None, "", []):
                existing[key] = value

        # Merge skills uniquely.
        existing_skills = existing.get("skills")
        incoming_skills = member.get("skills")
        if isinstance(existing_skills, list) or isinstance(incoming_skills, list):
            merged_skills = []
            seen_skills: set[str] = set()
            for skill in (existing_skills or []) + (incoming_skills or []):
                text = str(skill).strip()
                if not text:
                    continue
                key = text.lower()
                if key in seen_skills:
                    continue
                seen_skills.add(key)
                merged_skills.append(text)
            existing["skills"] = merged_skills

        # Merge engagements uniquely by engagement id.
        existing_engagements = existing.get("engagements") or []
        incoming_engagements = member.get("engagements") or []
        merged_engagements: list[dict[str, Any]] = []
        seen_eng_ids: set[str] = set()
        for eng in list(existing_engagements) + list(incoming_engagements):
            eng_id = str(eng.get("id") or "")
            if eng_id and eng_id in seen_eng_ids:
                continue
            if eng_id:
                seen_eng_ids.add(eng_id)
            merged_engagements.append(eng)
        existing["engagements"] = merged_engagements

    deduped_result: list[dict[str, Any]] = []
    for key in deduped_order:
        member = deduped_by_name[key]
        engagements = member.get("engagements") or []

        total_engagement_pct = 0.0
        for eng in engagements:
            level = eng.get("engagement_percentage")
            if level is None:
                level = eng.get("engagement_level")
            try:
                total_engagement_pct += float(level)
            except (ValueError, TypeError):
                pass

        member["engagement_pct"] = round(total_engagement_pct, 2)
        member["total_engagement_hours"] = sum(eng.get("engagement_hours", 0) for eng in engagements)
        member["total_tasks_completed"] = sum(eng.get("task_completed", 0) for eng in engagements)
        member["total_tasks_pending"] = sum(eng.get("task_pending", 0) for eng in engagements)
        member["projects_assigned"] = len(engagements)

        deduped_result.append(member)

    return deduped_result


# ── Resource Search & Auto-Allocation ────────────────────────────────────────


def _normalize_skills(raw: Any) -> list[str]:
    if not raw:
        return []
    if isinstance(raw, list):
        return [str(s).strip() for s in raw if str(s).strip()]
    if isinstance(raw, str):
        return [s.strip() for s in raw.split(",") if s.strip()]
    return []


def _score_skills(required: list[str], member_skills: list[str]) -> dict[str, Any]:
    if not required:
        return {"score": 0, "matched": [], "partial": []}
    lower_skills = [s.lower() for s in member_skills]
    matched: list[str] = []
    partial: list[str] = []
    for req in required:
        r = req.lower().strip()
        if any(s == r or r in s or s in r for s in lower_skills):
            matched.append(req)
        elif any(
            w
            for s in lower_skills
            for w in s.split()
            if len(w) > 3 and w in r
        ):
            partial.append(req)
    score = round(((len(matched) + len(partial) * 0.5) / len(required)) * 100)
    return {"score": score, "matched": matched, "partial": partial}


def _effective_skills(member: dict) -> list[str]:
    """
    Return the member's skills for scoring.
    When no explicit skills exist, derive tokens from the role so that
    e.g. "Senior Python Engineer" still matches "Python" and "Senior".
    Min 4 chars to avoid short abbreviations like "Dev" matching "Developer".
    """
    skills = _normalize_skills(member.get("skills") or member.get("Skills"))
    if skills:
        return skills
    role = (member.get("role") or "").replace("-", " ").replace("/", " ")
    return [w.strip() for w in role.split() if len(w.strip()) > 3]


def _score_role_match(required_role: str, member_role: str) -> int:
    """Score how well a member's role title matches the required role (0-100).
    Min 4-char tokens to avoid short abbreviations creating false matches.
    """
    if not required_role or not member_role:
        return 0
    req_tokens = [w.lower() for w in required_role.replace("-", " ").replace("/", " ").split() if len(w) > 3]
    mem_tokens = [w.lower() for w in member_role.replace("-", " ").replace("/", " ").split() if len(w) > 3]
    if not req_tokens or not mem_tokens:
        return 0
    matched = sum(1 for rt in req_tokens if any(rt == mt or (len(rt) >= 5 and len(mt) >= 5 and (rt in mt or mt in rt)) for mt in mem_tokens))
    return round(matched / len(req_tokens) * 100)


_llm_role_cache: dict[str, list[int]] = {}

_log = logging.getLogger(__name__)

# Debug helpers (opt-in via env var) -------------------------------------------------

def _env_truthy(name: str) -> bool:
    return (os.getenv(name) or "").strip().lower() in {"1", "true", "yes", "y", "on"}


def _env_truthy_default(name: str, default: bool) -> bool:
    raw = os.getenv(name)
    if raw is None:
        return default
    return raw.strip().lower() in {"1", "true", "yes", "y", "on"}


def _env_int(name: str, default: int) -> int:
    raw = (os.getenv(name) or "").strip()
    if not raw:
        return default
    try:
        return int(raw)
    except ValueError:
        return default


def _truncate_str(value: str) -> str:
    limit = max(20, _env_int("J2W_LOG_DB_STR_LIMIT", 500))
    if len(value) <= limit:
        return value
    return f"{value[:limit]}… (truncated, len={len(value)})"


def _sanitize_for_log(value: Any) -> Any:
    """
    Best-effort sanitizer to avoid leaking obvious secrets in logs.
    """
    secretish_keys = {"password", "passwd", "secret", "token", "api_key", "apikey", "key"}

    try:
        from bson import ObjectId  # type: ignore
    except Exception:  # pragma: no cover
        ObjectId = None  # type: ignore

    if value is None:
        return None
    if isinstance(value, str):
        return _truncate_str(value)
    if isinstance(value, (int, float, bool)):
        return value
    if isinstance(value, datetime):
        return value.isoformat()
    if ObjectId is not None and isinstance(value, ObjectId):  # type: ignore[arg-type]
        return str(value)
    if isinstance(value, list):
        return [_sanitize_for_log(v) for v in value]
    if isinstance(value, tuple):
        return [_sanitize_for_log(v) for v in list(value)]
    if isinstance(value, dict):
        sanitized: dict[str, Any] = {}
        for k, v in value.items():
            key = str(k)
            lower_key = key.strip().lower()
            if lower_key in secretish_keys or any(s in lower_key for s in secretish_keys):
                sanitized[key] = "***"
            else:
                sanitized[key] = _sanitize_for_log(v)
        if "_id" in sanitized:
            sanitized["_id"] = str(sanitized["_id"])
        return sanitized
    return str(value)


def _log_db_read(
    collection: str,
    *,
    query: dict[str, Any] | None = None,
    docs: Any,
    context: str | None = None,
) -> None:
    """
    Logs the DB collection name + data read for debugging.

    Enable with:
      - J2W_LOG_DB=1
    Optional:
      - J2W_LOG_DB_LIMIT=200
      - J2W_LOG_DB_PRETTY=1 (default: on when J2W_LOG_DB=1)
      - J2W_LOG_DB_STR_LIMIT=500
    """
    if not _env_truthy("J2W_LOG_DB"):
        return

    api_logger = logging.getLogger("delivery_tracker.api")
    limit = max(1, _env_int("J2W_LOG_DB_LIMIT", 200))
    pretty = _env_truthy_default("J2W_LOG_DB_PRETTY", True)

    if isinstance(docs, list):
        trimmed = docs[:limit]
        payload: dict[str, Any] = {
            "collection": collection,
            "query": query or {},
            "count": len(docs),
            "logged": len(trimmed),
            "truncated": len(docs) > len(trimmed),
            "docs": [_sanitize_for_log(d) for d in trimmed],
        }
    else:
        payload = {
            "collection": collection,
            "query": query or {},
            "doc": _sanitize_for_log(docs),
        }

    prefix = f"[DB READ] {context} " if context else "[DB READ] "
    rendered = json.dumps(
        payload,
        ensure_ascii=False,
        default=str,
        indent=2 if pretty else None,
    )
    if pretty:
        api_logger.info("%s\n%s", prefix.rstrip(), rendered)
    else:
        api_logger.info("%s%s", prefix, rendered)

# Keyword sets for job-family classification used as a safety floor after LLM scoring.
# When LLM gives 0 to a role that clearly belongs to the same job family as the required
# role, we apply a minimum score of 20 — same-family roles are always worth showing.
_FAMILY_KEYWORDS: dict[str, set[str]] = {
    # devops checked BEFORE developer so "devops" doesn't match "developer" keywords
    "devops": {
        "devops", "infrastructure", "cloud", "sysadmin", "platform",
        "site reliability", "sre", "kubernetes", "docker", "ci/cd",
        "devsecops",
    },
    "developer": {
        "developer", "engineer", "programmer", "coder",
        "fullstack", "full-stack", "full stack",
        "backend", "frontend", "front-end", "back-end",
        "software", "application", "web developer", "mobile developer",
        "android", "ios developer",
        "python", "java developer", "javascript", "typescript", "react developer",
        "golang", "rust developer", "scala", "kotlin", "swift developer",
    },
    "manager": {
        "manager", "programme", "program manager", "project manager", "delivery manager",
        "scrum", "agile", "product owner", "pod lead",
        "director", "head of", "vp of", "chief",
    },
    "data": {
        "data scientist", "data analyst", "data engineer", "analytics",
        "machine learning", "ml engineer", "ai engineer", "nlp", "etl",
        "business intelligence",
    },
    "qa": {
        "qa", "quality assurance", "tester", "test engineer", "sdet",
        "automation test",
    },
    "design": {
        "designer", "ux", "ui designer", "product design", "visual designer",
    },
    "seo": {
        "seo", "sem", "search engine", "content strategist", "marketing",
    },
}


def _role_family(role: str) -> str | None:
    lower = role.lower()
    for family, keywords in _FAMILY_KEYWORDS.items():
        if any(kw in lower for kw in keywords):
            return family
    return None


def _llm_score_roles_batch(required_role: str, member_roles: list[str]) -> list[int]:
    """
    Use Gemini to semantically score how well each member role matches required_role.
    Returns a list of scores 0–100 in the same order as member_roles.
    Falls back to _score_role_match on any failure.
    """
    if not required_role or not member_roles:
        return [_score_role_match(required_role, r) for r in member_roles]

    cache_key = required_role + "|" + "||".join(member_roles)
    if cache_key in _llm_role_cache:
        return _llm_role_cache[cache_key]

    api_key = os.getenv("GEMINI_API_KEY", "").strip()
    if not api_key:
        return [_score_role_match(required_role, r) for r in member_roles]

    roles_numbered = "\n".join(f"{i + 1}. {role}" for i, role in enumerate(member_roles))
    prompt = (
        f'You are matching team members to a project resource slot.\n\n'
        f'Required role: "{required_role}"\n\n'
        f"First, identify the CORE JOB FUNCTION of the required role "
        f"(e.g. 'writes application code', 'manages project delivery', 'handles infra/deployment', 'data/analytics', 'design/UX', 'SEO/marketing').\n\n"
        f"Then score each member role 0–100 based on whether they share that same core function:\n"
        f"- 80–100: Same core function, same or very similar specialization\n"
        f"- 50–79:  Same core function, adjacent specialization (e.g. Project Manager ↔ Program Manager)\n"
        f"- 20–49:  Same core function, different technology/stack/language — IMPORTANT: if the core function "
        f"is 'writes application code', then Software Engineer, Java Developer, Full Stack Dev, Backend Dev, "
        f"Mobile Dev, Web Developer all share this function and must score AT LEAST 20, regardless of language\n"
        f"- 1–19:   Overlapping function (e.g. Architect touches code but is not primarily a developer)\n"
        f"- 0:      Completely different core function — ONLY for roles with ZERO functional overlap "
        f"(DevOps/Infra roles for PM or coding positions; SEO/Marketing for any tech role; HR/Finance for tech roles)\n\n"
        f"Member roles:\n{roles_numbered}\n\n"
        f"Return ONLY a JSON array of integers (one per role, same order). Nothing else."
    )

    try:
        with httpx.Client(timeout=20.0) as client:
            resp = client.post(
                "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent",
                params={"key": api_key},
                json={
                    "contents": [{"role": "user", "parts": [{"text": prompt}]}],
                    "generationConfig": {"temperature": 0.1, "maxOutputTokens": 512},
                },
            )
            resp.raise_for_status()
            text = resp.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
            m = re.search(r"\[[\d,\s]+\]", text)
            if m:
                scores = json.loads(m.group())
                if len(scores) == len(member_roles):
                    result = [min(100, max(0, int(s))) for s in scores]
                    # Safety floor: if the LLM gave 0 to a role that clearly belongs
                    # to the same job family as the required role, bump it to 20.
                    # This handles the model's bias of treating "Software Engineer"
                    # and "Software Developer" as different families.
                    req_family = _role_family(required_role)
                    if req_family:
                        result = [
                            max(score, 20) if score == 0 and _role_family(member_roles[i]) == req_family else score
                            for i, score in enumerate(result)
                        ]
                    _llm_role_cache[cache_key] = result
                    return result
    except Exception as exc:
        _log.warning("LLM role scoring failed, falling back to token match: %s", exc)

    return [_score_role_match(required_role, r) for r in member_roles]


def _build_project_timeline_lookup(database: Database) -> dict[str, dict[str, Any]]:
    """Build project_id → {name, start_date, end_date} from projects + milestones."""
    projects_by_id: dict[str, dict] = {}
    for doc in database["projects"].find({}):
        pid = doc.get("id", "")
        if pid:
            projects_by_id[pid] = {"name": _resolve_project_display_name(doc, str(pid)), "starts": [], "ends": []}

    for m in database["milestones"].find({}):
        pid = m.get("project_id", "")
        if pid not in projects_by_id:
            continue
        # Support both camelCase (from proposal save) and snake_case (from health tracker)
        start = (
            m.get("actual_start")
            or m.get("plannedStart")
            or m.get("planned_start")
        )
        end = (
            m.get("actual_end_eta")
            or m.get("plannedEnd")
            or m.get("planned_end")
        )
        if start:
            projects_by_id[pid]["starts"].append(str(start)[:10])
        if end:
            projects_by_id[pid]["ends"].append(str(end)[:10])

    lookup: dict[str, dict[str, Any]] = {}
    for pid, data in projects_by_id.items():
        starts = sorted(data["starts"])
        ends = sorted(data["ends"])
        lookup[pid] = {
            "name": data["name"],
            "start_date": starts[0] if starts else None,
            "end_date": ends[-1] if ends else None,
        }
    return lookup


def _build_engagement_detail(
    database: Database,
    project_lookup: dict[str, dict[str, Any]],
) -> dict[str, list[dict[str, Any]]]:
    """Build member_id → list of project allocation details."""
    detail: dict[str, list[dict[str, Any]]] = {}
    for eng in database["team_members_engagement"].find({}):
        member_id = eng.get("team_member_id", "")
        if not member_id:
            continue
        project_id = eng.get("project_id", "")
        raw = eng.get("engagement_level") or eng.get("engagement_percentage") or 0
        try:
            level = float(raw)
        except (ValueError, TypeError):
            level = 0.0
        proj = project_lookup.get(project_id, {})
        detail.setdefault(member_id, []).append({
            "project_id": project_id,
            "project_name": proj.get("name") or project_id,
            "engagement_level": level,
            "start_date": proj.get("start_date"),
            "end_date": proj.get("end_date"),
        })
    return detail


def _build_engagement_map(database: Database) -> dict[str, float]:
    eng_map: dict[str, float] = {}
    for eng in database["team_members_engagement"].find({}):
        member_id = eng.get("team_member_id", "")
        if not member_id:
            continue
        raw = eng.get("engagement_level") or eng.get("engagement_percentage") or 0
        try:
            level = float(raw)
        except (ValueError, TypeError):
            level = 0.0
        eng_map[member_id] = eng_map.get(member_id, 0.0) + level
    return eng_map


def _filter_members_by_resource_type(
    members: list[dict],
    resource_type: str | None,
) -> list[dict]:
    """
    Filter members by resource_type.
    - "Internal": resource_type == "Internal" OR member_type == "Consultant"
    - "External": resource_type == "External"
    - None / anything else: no filter (return all)
    """
    if not resource_type:
        return members
    rt = resource_type.strip().lower()
    if rt == "internal":
        return [
            m for m in members
            if (m.get("resource_type") or "").strip().lower() == "internal"
            or (m.get("member_type") or "").strip().lower() == "consultant"
        ]
    if rt == "external":
        return [
            m for m in members
            if (m.get("resource_type") or "").strip().lower() == "external"
        ]
    return members


def _resolve_resource_type(member: dict) -> str:
    """Derive the display resource type label for a member."""
    if (member.get("member_type") or "").strip().lower() == "consultant":
        return "Consultant"
    rt = (member.get("resource_type") or "").strip()
    return rt if rt else "Internal"


def search_resources(
    database: Database,
    required_skills: list[str],
    bandwidth_needed: int = 0,
    **_: Any,
) -> dict[str, Any]:
    """
    Search team members matching required skills and availability.
    Composite score = 70% skill match + 30% bandwidth availability.
    Skills are matched strictly against the member's skills array.
    """
    raw_members = list(database["team_members"].find({"is_active": {"$ne": False}}))
    members = [to_plain_document(doc) for doc in raw_members if doc]

    project_lookup = _build_project_timeline_lookup(database)
    eng_detail = _build_engagement_detail(database, project_lookup)

    results: list[dict[str, Any]] = []
    for member in members:
        member_id = member.get("id", "")
        projects = eng_detail.get(member_id, [])
        committed = sum(p["engagement_level"] for p in projects)
        avail_bw = round(max(0.0, 100.0 - committed), 1)

        # Skip members who don't have enough bandwidth when a minimum is specified
        if bandwidth_needed and avail_bw < bandwidth_needed:
            continue

        # Match strictly against the skills array — no role-token fallback
        member_skills = _normalize_skills(member.get("skills") or member.get("Skills"))
        skill_result = _score_skills(required_skills, member_skills)

        # Composite: 70% skill match + 30% bandwidth availability
        composite = round(skill_result["score"] * 0.7 + avail_bw * 0.3)

        results.append({
            "id": member_id,
            "name": member.get("name", ""),
            "role": member.get("role", ""),
            "member_type": member.get("member_type", ""),
            "resource_type": _resolve_resource_type(member),
            "initials": member.get("initials", ""),
            "color_hex": member.get("color_hex", "#6366f1"),
            "skills": member_skills,
            "available_bandwidth": avail_bw,
            "committed_bandwidth": round(committed, 1),
            "skill_score": skill_result["score"],
            "matched_skills": skill_result["matched"],
            "partial_skills": skill_result["partial"],
            "composite_score": composite,
            "active_projects": projects,
        })

    results.sort(key=lambda x: (-x["composite_score"], -x["available_bandwidth"]))

    return {"members": results, "total": len(results)}


def auto_allocate_resources(
    database: Database,
    resources_input: list[dict[str, Any]],
    project_start: str | None = None,
    project_end: str | None = None,
) -> list[dict[str, Any]]:
    """
    Auto-allocate best team members to resource slots.
    Each member assigned at most once; minimum score 50, minimum availability 20%.
    """
    raw_members = list(database["team_members"].find({"is_active": {"$ne": False}}))
    members = [to_plain_document(doc) for doc in raw_members if doc]

    eng_map = _build_engagement_map(database)
    used_bw: dict[str, float] = {}
    assigned_ids: set[str] = set()

    results: list[dict[str, Any]] = []
    for res in resources_input:
        required_skills = res.get("skills", [])
        bandwidth_needed = int(res.get("bandwidth", 100))
        assignments: list[dict[str, Any]] = []
        open_bw = float(bandwidth_needed)

        candidates: list[tuple[dict, dict, float]] = []
        for member in members:
            member_id = member.get("id", "")
            if member_id in assigned_ids:
                continue
            committed = eng_map.get(member_id, 0.0) + used_bw.get(member_id, 0.0)
            avail = max(0.0, 100.0 - committed)
            if avail < 20.0:
                continue
            effective = _effective_skills(member)
            skill_result = _score_skills(required_skills, effective)
            # Accept candidates with any skill match, or if no required skills specified
            if required_skills and skill_result["score"] < 10:
                continue
            candidates.append((member, skill_result, avail))

        candidates.sort(key=lambda x: (-x[1]["score"], -x[2]))

        if candidates:
            best_member, best_skill, best_avail = candidates[0]
            alloc_bw = min(open_bw, best_avail)
            member_id = best_member.get("id", "")
            assignments.append({
                "type": "internal",
                "id": member_id,
                "name": best_member.get("name", ""),
                "role": best_member.get("role", ""),
                "bw": round(alloc_bw, 1),
                "score": best_skill["score"],
            })
            used_bw[member_id] = used_bw.get(member_id, 0.0) + alloc_bw
            assigned_ids.add(member_id)
            open_bw -= alloc_bw

        tbd_bw = max(0.0, open_bw)
        results.append({
            "resourceId": res.get("id", ""),
            "role": res.get("role", ""),
            "requiredBW": bandwidth_needed,
            "assignments": assignments,
            "tbdBW": round(tbd_bw, 1),
            "status": "matched" if tbd_bw <= 0 else "partial",
        })

    return results


def list_records(database: Database, table: str, query_params: dict[str, str]) -> list[dict[str, Any]]:
    ensure_table(table)
    cursor = database[table].find(build_filter(query_params))

    order_by = query_params.get("orderBy")
    if order_by:
        descending = query_params.get("ascending") == "false"
        cursor = cursor.sort(order_by, -1 if descending else 1)

    limit = query_params.get("limit")
    if limit is not None:
        try:
            limit_value = int(limit)
        except ValueError:
            limit_value = 0
        if limit_value > 0:
            cursor = cursor.limit(limit_value)

    offset = query_params.get("offset")
    if offset is not None:
        try:
            offset_value = int(offset)
        except ValueError:
            offset_value = 0
        if offset_value > 0:
            cursor = cursor.skip(offset_value)

    return [to_plain_document(doc) or {} for doc in cursor]


def get_record(database: Database, table: str, record_id: str) -> dict[str, Any]:
    ensure_table(table)
    document = database[table].find_one({"id": str(record_id)})
    if document is None:
        raise HTTPException(status_code=404, detail=f"{table} record not found")
    return to_plain_document(document) or {}


def create_records(database: Database, table: str, payload: Any) -> list[dict[str, Any]] | dict[str, Any]:
    ensure_table(table)

    if not isinstance(payload, (dict, list)):
        raise HTTPException(status_code=400, detail="Invalid body")

    if isinstance(payload, list):
        documents = [normalize_on_insert(table, item) for item in payload]
        if not documents:
            return []
        database[table].insert_many(documents)

        # Regenerate week data for milestones
        if table == "milestones":
            for doc in documents:
                project_id = doc.get("project_id")
                if project_id:
                    regenerate_project_milestone_weeks(database, project_id)

        return [to_plain_document(document) or {} for document in documents]

    document = normalize_on_insert(table, payload)
    database[table].insert_one(document)

    # Regenerate week data for milestones
    if table == "milestones":
        project_id = document.get("project_id")
        if project_id:
            regenerate_project_milestone_weeks(database, project_id)

    return to_plain_document(document) or {}


def update_milestone_health(database: Database, milestone_id: str, milestone_type: str, status: str, date: str | None = None) -> dict[str, Any] | None:
    """
    Update a specific milestone health type (practice, signoff, or invoice).

    Args:
        database: MongoDB database connection
        milestone_id: ID of the milestone to update
        milestone_type: Type of update - "practice", "signoff", or "invoice"
        status: New status value
        date: Optional date for signoff/invoice (when marked as "Done")

    Returns:
        Updated milestone document

    Example:
        # Update practice status
        update_milestone_health(db, "m123", "practice", "At Risk")

        # Mark signoff as done with date
        update_milestone_health(db, "m123", "signoff", "Done", "2026-02-07")
    """
    if milestone_type not in {"practice", "signoff", "invoice"}:
        raise HTTPException(status_code=400, detail=f"Invalid milestone type: {milestone_type}")

    # Validate status values
    if milestone_type == "practice":
        valid_statuses = {"On Track", "At Risk", "Blocked", "Completed"}
        if status not in valid_statuses:
            raise HTTPException(status_code=400, detail=f"Invalid practice status: {status}")
        update_dict = {"status": status}

    elif milestone_type == "signoff":
        valid_statuses = {"Done", "Pending"}
        if status not in valid_statuses:
            raise HTTPException(status_code=400, detail=f"Invalid signoff status: {status}")
        update_dict = {"client_signoff_status": status}
        if status == "Done" and date:
            update_dict["signedoff_date"] = date

    elif milestone_type == "invoice":
        valid_statuses = {"Done", "Pending"}
        if status not in valid_statuses:
            raise HTTPException(status_code=400, detail=f"Invalid invoice status: {status}")
        update_dict = {"invoice_status": status}
        if status == "Done" and date:
            update_dict["invoice_raised_date"] = date

    # Perform the update
    return patch_record(database, "milestones", milestone_id, update_dict)



def update_week_status(
    database: Database,
    milestone_id: str,
    milestone_type: str,
    week_number: int,
    payload: dict[str, Any]
) -> dict[str, Any] | None:
    """
    Update status for a specific week in a milestone.

    If the week exists, it will be updated.
    If the week doesn't exist, it will be created and added to the array.
    All other weeks remain unchanged.

    Args:
        database: MongoDB database connection
        milestone_id: ID of the milestone to update
        milestone_type: Type of update - "practice", "signoff", or "invoice"
        week_number: The week number to update or create
        payload: Week update payload containing:
            {
                "week_status": "At Risk",        // New status (required)
                "week_label": "Feb 16-22, 2026", // Display label (required)
                "color": "orange",               // Color indicator (required)
                "date": "2026-02-20"             // Week date (required)
            }

    Returns:
        Updated milestone document with the week added/updated

    Example:
        # Update existing week 2
        update_week_status(
            db, "m1", "practice", 2,
            {
                "week_status": "At Risk",
                "week_label": "Feb 16-22, 2026",
                "color": "orange",
                "date": "2026-02-20"
            }
        )

        # Create new week 5 if it doesn't exist
        update_week_status(
            db, "m1", "practice", 5,
            {
                "week_status": "On Track",
                "week_label": "Mar 9-15, 2026",
                "color": "green",
                "date": "2026-03-09"
            }
        )
    """
    # Validate milestone type
    if milestone_type not in {"practice", "signoff", "invoice"}:
        raise HTTPException(status_code=400, detail="Invalid milestone type")

    # Get the milestone
    milestone = database["milestones"].find_one({"id": milestone_id})
    if not milestone:
        raise HTTPException(status_code=404, detail="Milestone not found")

    # Get the appropriate weeks array key
    if milestone_type == "practice":
        weeks_key = "practice_weeks"
    elif milestone_type == "signoff":
        weeks_key = "signoff_weeks"
    else:  # invoice
        weeks_key = "invoice_weeks"

    # Get the weeks array
    weeks_array = milestone.get(weeks_key, [])

    # Find and update the week, or create it if it doesn't exist
    updated = False
    for i, week in enumerate(weeks_array):
        if week.get("week_number") == week_number:
            # Update existing week with new data
            weeks_array[i]["status"] = payload.get("week_status", week.get("status"))
            weeks_array[i]["color"] = payload.get("color", week.get("color"))
            weeks_array[i]["week_label"] = payload.get("week_label", week.get("week_label"))
            weeks_array[i]["date"] = payload.get("date", week.get("date"))
            updated = True
            break

    if not updated:
        # Create new week if it doesn't exist
        new_week = {
            "week_number": week_number,
            "status": payload.get("week_status"),
            "color": payload.get("color"),
            "week_label": payload.get("week_label"),
            "date": payload.get("date")
        }
        weeks_array.append(new_week)
        # Sort by week_number to maintain chronological order
        weeks_array.sort(key=lambda w: w.get("week_number", 0))

    # Save updated weeks array
    database["milestones"].update_one(
        {"id": milestone_id},
        {"$set": {weeks_key: weeks_array, "updated_at": utc_now_iso()}}
    )

    # Return updated milestone
    updated_milestone = database["milestones"].find_one({"id": milestone_id})
    return to_plain_document(updated_milestone)


def delete_week_status(
    database: Database,
    milestone_id: str,
    milestone_type: str,
    week_number: int,
) -> dict[str, Any] | None:
    """Remove a specific week entry from a milestone's weeks array."""
    if milestone_type not in {"practice", "signoff", "invoice"}:
        raise HTTPException(status_code=400, detail="Invalid milestone type")

    milestone = database["milestones"].find_one({"id": milestone_id})
    if not milestone:
        raise HTTPException(status_code=404, detail="Milestone not found")

    weeks_key = f"{milestone_type}_weeks"
    weeks_array = milestone.get(weeks_key, [])
    filtered = [w for w in weeks_array if w.get("week_number") != week_number]

    if len(filtered) == len(weeks_array):
        raise HTTPException(status_code=404, detail=f"Week {week_number} not found")

    database["milestones"].update_one(
        {"id": milestone_id},
        {"$set": {weeks_key: filtered, "updated_at": utc_now_iso()}}
    )

    updated = database["milestones"].find_one({"id": milestone_id})
    return to_plain_document(updated)


def patch_record(database: Database, table: str, record_id: str, changes: Any) -> dict[str, Any] | None:
    ensure_table(table)

    if not isinstance(changes, dict):
        raise HTTPException(status_code=400, detail="Invalid body")

    next_changes = deepcopy(changes)
    if table in {"projects", "milestones", "team_members"}:
        next_changes["updated_at"] = utc_now_iso()

    database[table].update_one({"id": str(record_id)}, {"$set": next_changes})
    updated = database[table].find_one({"id": str(record_id)})

    # Regenerate week data for milestones if relevant fields changed
    if table == "milestones" and updated:
        project_id = updated.get("project_id")
        # Check if any week-generating fields were modified
        week_fields = {"actual_start", "actual_end_eta", "status", "client_signoff_status",
                      "signedoff_date", "invoice_status", "invoice_raised_date"}
        if project_id and any(field in next_changes for field in week_fields):
            regenerate_project_milestone_weeks(database, project_id)

    return to_plain_document(updated)


def replace_record(database: Database, table: str, record_id: str, replacement: Any) -> dict[str, Any]:
    ensure_table(table)

    if not isinstance(replacement, dict):
        raise HTTPException(status_code=400, detail="Invalid body")

    existing = database[table].find_one({"id": str(record_id)})
    if existing is None:
        raise HTTPException(status_code=404, detail=f"{table} record not found")

    next_document = {**existing, **replacement, "id": str(record_id)}
    if table in {"projects", "milestones", "team_members"}:
        next_document["updated_at"] = utc_now_iso()

    database[table].replace_one({"id": str(record_id)}, next_document)
    updated = database[table].find_one({"id": str(record_id)})

    # Regenerate week data for milestones
    if table == "milestones" and updated:
        project_id = updated.get("project_id")
        if project_id:
            regenerate_project_milestone_weeks(database, project_id)

    return to_plain_document(updated) or {}


def delete_record(database: Database, table: str, record_id: str) -> dict[str, Any]:
    ensure_table(table)

    existing = database[table].find_one({"id": str(record_id)})

    if table == "project_documents" and existing and existing.get("path"):
        full_path = Path(UPLOADS_DIR.parent, existing["path"])
        if full_path.exists():
            full_path.unlink()

    database[table].delete_one({"id": str(record_id)})
    return {"ok": True, "deleted": to_plain_document(existing)}


def save_upload(
    database: Database,
    file: UploadFile,
    project_id: str,
    update_id: str | None = None,
    category: str | None = None,
) -> dict[str, Any]:
    if not project_id:
        raise HTTPException(status_code=400, detail="Project ID is required")

    UPLOADS_DIR.mkdir(parents=True, exist_ok=True)

    safe_name = Path(file.filename or "upload.bin").name
    unique_name = f"{int(datetime.now(timezone.utc).timestamp() * 1000)}-{safe_name}"
    disk_path = UPLOADS_DIR / unique_name

    with disk_path.open("wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    file_details: dict[str, Any] = {
        "name": safe_name,
        "size": disk_path.stat().st_size,
        "type": file.content_type or "application/octet-stream",
        "path": f"uploads/{unique_name}",
        "project_id": project_id,
    }
    if category:
        file_details["category"] = category

    document = normalize_on_insert("project_documents", file_details)
    database["project_documents"].insert_one(document)

    if update_id:
        database["project_updates"].update_one(
            {"id": update_id},
            {"$set": {"file_path": file_details["path"], "file_name": file_details["name"]}},
        )

    return to_plain_document(document) or {}


CLOUD_UPLOAD_FOLDER = "centralized_delivery_tracker_records"


async def upload_to_cloud(
    database: Database,
    file: UploadFile,
    settings: Any,
    project_id: str,
    update_id: str | None = None,
    category: str | None = None,
) -> dict[str, Any]:
    from azure.storage.blob import BlobServiceClient, ContentSettings

    if not project_id:
        raise HTTPException(status_code=400, detail="project_id is required")

    account_name = settings.azure_storage_account_name
    account_key = settings.azure_storage_account_key
    container_name = settings.azure_storage_container_name

    if not account_name or not account_key or not container_name:
        raise HTTPException(
            status_code=500,
            detail="Azure Storage is not configured. Set AZURE_STORAGE_ACCOUNT_NAME, AZURE_STORAGE_ACCOUNT_KEY, and AZURE_STORAGE_CONTAINER_NAME.",
        )

    safe_name = Path(file.filename or "upload.bin").name
    timestamp = int(datetime.now(timezone.utc).timestamp() * 1000)
    blob_name = f"{CLOUD_UPLOAD_FOLDER}/{timestamp}-{safe_name}"

    connection_string = (
        f"DefaultEndpointsProtocol=https;"
        f"AccountName={account_name};"
        f"AccountKey={account_key};"
        f"EndpointSuffix=core.windows.net"
    )

    content = await file.read()

    try:
        blob_client = BlobServiceClient.from_connection_string(connection_string).get_blob_client(
            container=container_name,
            blob=blob_name,
        )
        blob_client.upload_blob(
            content,
            overwrite=True,
            content_settings=ContentSettings(content_type=file.content_type or "application/octet-stream"),
        )
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Azure Blob upload failed: {exc}") from exc

    blob_url = f"https://{account_name}.blob.core.windows.net/{container_name}/{blob_name}"

    file_details: dict[str, Any] = {
        "name": safe_name,
        "size": len(content),
        "type": file.content_type or "application/octet-stream",
        "path": blob_url,
        "blob_name": blob_name,
        "project_id": project_id,
        "storage": "azure_blob",
    }
    if category:
        file_details["category"] = category

    query: dict[str, Any] = {"project_id": project_id}
    if category:
        query["category"] = category
    existing_doc = database["project_documents"].find_one(query)

    if existing_doc:
        # Find next free path slot: path → path2 → path3 → ...
        path_key = "path"
        counter = 2
        while existing_doc.get(path_key):
            path_key = f"path{counter}"
            counter += 1
        name_key = "name" if path_key == "path" else f"name{path_key[4:]}"
        database["project_documents"].update_one(
            {"id": existing_doc["id"]},
            {"$set": {path_key: blob_url, name_key: safe_name}},
        )
        result_doc = to_plain_document(
            database["project_documents"].find_one({"id": existing_doc["id"]})
        ) or {}
    else:
        document = normalize_on_insert("project_documents", file_details)
        database["project_documents"].insert_one(document)
        result_doc = to_plain_document(document) or {}

    if update_id:
        existing_update = database["project_updates"].find_one({"id": update_id})
        if existing_update and existing_update.get("file_path"):
            # Find next free path slot: file_path → file_path2 → file_path3 → ...
            path_key = "file_path"
            counter = 2
            while existing_update.get(path_key):
                path_key = f"file_path{counter}"
                counter += 1
            name_key = "file_name" if path_key == "file_path" else f"file_name{path_key[9:]}"
            database["project_updates"].update_one(
                {"id": update_id},
                {"$set": {path_key: blob_url, name_key: safe_name}},
            )
        else:
            database["project_updates"].update_one(
                {"id": update_id},
                {"$set": {"file_path": blob_url, "file_name": safe_name}},
            )

    return result_doc


async def get_skills_for_designation(designation: str, azure: Any) -> dict[str, Any]:
    from .ai.azure_openai_service import AzureChatConfig
    from .ai.prompts import DESIGNATION_SKILLS_SYSTEM, DESIGNATION_SKILLS_USER

    designation = designation.strip()
    if not designation:
        raise HTTPException(status_code=400, detail="Designation is required")

    prompt = DESIGNATION_SKILLS_USER.format(designation=designation)
    config = AzureChatConfig(temperature=0.2, max_tokens=2048)

    result = await azure.generate_json(
        prompt=prompt,
        system_prompt=DESIGNATION_SKILLS_SYSTEM,
        config=config,
    )

    if not result or not isinstance(result, dict):
        raise HTTPException(status_code=502, detail="LLM returned an unexpected response")

    return result


def _normalize_reporting_window(
    start_date_str: str | None,
    end_date_str: str | None,
) -> tuple[datetime, datetime]:
    now = datetime.now(timezone.utc)
    month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    if now.month == 12:
        next_month = now.replace(year=now.year + 1, month=1, day=1)
    else:
        next_month = now.replace(month=now.month + 1, day=1)
    month_end = next_month.replace(hour=23, minute=59, second=59, microsecond=999999) - timedelta(days=1)

    parsed_start = parse_date(start_date_str) if start_date_str else month_start
    parsed_end = parse_date(end_date_str) if end_date_str else month_end

    if not parsed_start:
        parsed_start = month_start
    if not parsed_end:
        parsed_end = month_end

    parsed_start = parsed_start.replace(hour=0, minute=0, second=0, microsecond=0)
    parsed_end = parsed_end.replace(hour=23, minute=59, second=59, microsecond=999999)

    if parsed_start > parsed_end:
        raise HTTPException(status_code=400, detail="start_date must be less than or equal to end_date")

    return parsed_start, parsed_end


def _window_intersection(
    a_start: datetime,
    a_end: datetime,
    b_start: datetime,
    b_end: datetime,
) -> tuple[datetime, datetime] | None:
    start = max(a_start, b_start)
    end = min(a_end, b_end)
    if start > end:
        return None
    return start, end


def _extract_update_datetime(update: dict[str, Any]) -> datetime | None:
    return parse_date(update.get("activity_date")) or parse_date(update.get("created_at"))


def _extract_update_text(update: dict[str, Any]) -> str | None:
    value = update.get("content") or update.get("comment")
    if not value:
        return None
    text = str(value).strip()
    return text or None


def _safe_filename_part(value: str | None) -> str:
    raw = (value or "report").strip()
    cleaned = re.sub(r"[^a-zA-Z0-9._-]+", "-", raw).strip("-._")
    return cleaned[:80] or "report"


def _collect_project_updates(
    database: Database,
    project_id: str,
    window_start: datetime,
    window_end: datetime,
) -> list[dict[str, Any]]:
    start_date = window_start.date().isoformat()
    end_date = window_end.date().isoformat()
    start_ts = window_start.astimezone(timezone.utc).isoformat().replace("+00:00", "Z")
    end_ts = window_end.astimezone(timezone.utc).isoformat().replace("+00:00", "Z")

    # DB-side windowing so we don't fetch all project updates when the caller only
    # needs a bounded date range. Most records use `activity_date` (YYYY-MM-DD).
    # For older/edge records that don't have activity_date, fall back to created_at.
    query: dict[str, Any] = {
        "project_id": project_id,
        "$or": [
            {"activity_date": {"$gte": start_date, "$lte": end_date}},
            {
                "activity_date": {"$exists": False},
                "created_at": {"$gte": start_ts, "$lte": end_ts},
            },
            {
                "activity_date": None,
                "created_at": {"$gte": start_ts, "$lte": end_ts},
            },
            {
                "activity_date": "",
                "created_at": {"$gte": start_ts, "$lte": end_ts},
            },
        ],
    }
    updates = list(database["project_updates"].find(query))
    _log_db_read(
        "project_updates",
        query=query,
        docs=updates,
        context=f"project_id={project_id} (db-windowed raw)",
    )
    filtered: list[dict[str, Any]] = []
    for update in updates:
        dt = _extract_update_datetime(update)
        if not dt or dt < window_start or dt > window_end:
            continue
        text = _extract_update_text(update)
        if not text:
            continue
        filtered.append(
            {
                "dt": dt,
                "text": text,
                "activity_date": update.get("activity_date"),
                "created_at": update.get("created_at"),
            }
        )
    filtered.sort(key=lambda item: item["dt"])
    _log_db_read(
        "project_updates",
        query={
            "project_id": project_id,
            "window_start": window_start.date().isoformat(),
            "window_end": window_end.date().isoformat(),
        },
        docs=filtered,
        context=f"project_id={project_id} (filtered)",
    )
    return filtered


async def generate_multi_project_report(
    database: Database,
    project_ids: list[str],
    start_date_str: str | None = None,
    end_date_str: str | None = None
) -> list[dict[str, Any]]:
    """
    Generate an AI-backed status report for multiple projects within a strict date window.
    Progress is bounded by:
    1) Global reporting window.
    2) Individual milestone start/end window.
    """
    report_start, report_end = _normalize_reporting_window(start_date_str, end_date_str)
    report_start_label = report_start.date().isoformat()
    report_end_label = report_end.date().isoformat()
    report_results = []

    if _env_truthy("J2W_LOG_DB"):
        logging.getLogger("delivery_tracker.api").info(
            "[REPORT] /reports/generate %s",
            json.dumps(
                {
                    "project_ids": list(dict.fromkeys(project_ids)),
                    "window_start": report_start_label,
                    "window_end": report_end_label,
                },
                ensure_ascii=False,
                default=str,
            ),
        )

    for pid in dict.fromkeys(project_ids):
        project_query = {"id": pid}
        project = database["projects"].find_one(project_query)
        _log_db_read("projects", query=project_query, docs=project, context=f"project_id={pid}")
        if not project:
            continue

        project_updates = _collect_project_updates(database, pid, report_start, report_end)
        milestones_query = {"project_id": pid}
        milestones = list(database["milestones"].find(milestones_query))
        milestones.sort(key=lambda m: (str(m.get("milestone_code") or ""), str(m.get("description") or "")))

        # Keep milestones aligned to the reporting window (inclusive) so both the
        # report output and debug logs reflect only in-window workstreams.
        in_window_milestones: list[dict[str, Any]] = []
        for m in milestones:
            milestone_start = (
                parse_date(m.get("actual_start"))
                or parse_date(m.get("planned_start"))
                or report_start
            )
            milestone_end = (
                parse_date(m.get("actual_end_eta"))
                or parse_date(m.get("planned_end"))
                or parse_date(m.get("planned_end_eta"))
                or report_end
            )
            if milestone_end < milestone_start:
                milestone_end = milestone_start
            if not _window_intersection(report_start, report_end, milestone_start, milestone_end):
                continue
            in_window_milestones.append(m)

        milestones = in_window_milestones
        _log_db_read(
            "milestones",
            query={
                **milestones_query,
                "window_start": report_start_label,
                "window_end": report_end_label,
            },
            docs=milestones,
            context=f"project_id={pid} (in-window)",
        )
        milestone_insights = []

        for m in milestones:
            m_id = m.get("id")
            m_name = m.get("milestone_code") or m.get("description", "Unknown")

            milestone_start = (
                parse_date(m.get("actual_start"))
                or parse_date(m.get("planned_start"))
                or report_start
            )
            milestone_end = (
                parse_date(m.get("actual_end_eta"))
                or parse_date(m.get("planned_end"))
                or parse_date(m.get("planned_end_eta"))
                or report_end
            )
            if milestone_end < milestone_start:
                milestone_end = milestone_start

            bounded_window = _window_intersection(report_start, report_end, milestone_start, milestone_end)
            if not bounded_window:
                continue
            bounded_start, bounded_end = bounded_window

            bounded_updates = [
                u for u in project_updates
                if bounded_start <= u["dt"] <= bounded_end
            ]
            logs = [f"{u['dt'].date().isoformat()}: {u['text']}" for u in bounded_updates]

            insight = await insight_service.generate_milestone_insight(
                m_name,
                m.get("description", ""),
                m.get("status", "Unknown"),
                logs,
                start_date=bounded_start.date().isoformat(),
                end_date=bounded_end.date().isoformat(),
                milestone_meta={
                    "milestone_code": m.get("milestone_code"),
                    "planned_start": m.get("planned_start"),
                    "planned_end": m.get("planned_end"),
                    "planned_end_eta": m.get("planned_end_eta"),
                    "actual_start": m.get("actual_start"),
                    "actual_end_eta": m.get("actual_end_eta"),
                    "actual_end": m.get("actual_end"),
                    "completion_pct": m.get("completion_pct"),
                    "invoice_status": m.get("invoice_status"),
                    "client_signoff_status": m.get("client_signoff_status"),
                    "blocker": m.get("blocker"),
                },
            )

            milestone_insights.append({
                "id": m_id,
                "name": m_name,
                "description": m.get("description"),
                "status": m.get("status"),
                "planned_start": m.get("planned_start"),
                "planned_end": m.get("planned_end") or m.get("planned_end_eta"),
                "actual_start": m.get("actual_start"),
                "actual_end": m.get("actual_end_eta"),
                "window_start": bounded_start.date().isoformat(),
                "window_end": bounded_end.date().isoformat(),
                "log_count": len(logs),
                "timeline_summary": insight,
            })

        proj_summary = await insight_service.generate_project_executive_summary(
            project.get("name", "Unnamed Project"),
            [{"name": m["name"], "summary": m["timeline_summary"]} for m in milestone_insights],
            project_updates=[f"{u['dt'].date().isoformat()}: {u['text']}" for u in project_updates],
            start_date=report_start_label,
            end_date=report_end_label,
        )

        report_results.append({
            "project_id": pid,
            "project_name": project.get("name"),
            "manager": project.get("delivery_manager"),
            "spoc": project.get("client_spoc"),
            "status": project.get("status"),
            "window_start": report_start_label,
            "window_end": report_end_label,
            "project_update_count_in_window": len(project_updates),
            "executive_summary": proj_summary,
            "milestones": milestone_insights,
        })

    return report_results


async def generate_project_insight_report(
    database: Database,
    project_ids: list[str],
    start_date_str: str | None = None,
    end_date_str: str | None = None,
) -> list[dict[str, Any]]:
    """
    Backward-compatible wrapper.
    """
    return await generate_multi_project_report(database, project_ids, start_date_str, end_date_str)


def _report_lines(report: dict[str, Any]) -> list[str]:
    lines: list[str] = []
    lines.append(f"Project Status Report: {report.get('project_name') or 'Unnamed Project'}")
    lines.append(f"Reporting Window: {report.get('window_start', '')} to {report.get('window_end', '')}")
    lines.append("")
    lines.append("Executive Summary")
    lines.append(str(report.get("executive_summary") or ""))
    lines.append("")
    lines.append("Milestone Timeline Stories")
    for idx, milestone in enumerate(report.get("milestones") or [], start=1):
        lines.append(
            f"{idx}. {milestone.get('name', 'Milestone')} | "
            f"Status: {milestone.get('status') or 'Unknown'} | "
            f"Window: {milestone.get('window_start', '')} to {milestone.get('window_end', '')}"
        )
        lines.append(str(milestone.get("timeline_summary") or "No in-window activity."))
        lines.append("")
    return lines


def _resolve_j2w_logo_path() -> Path | None:
    service_file = Path(__file__).resolve()
    repo_root = service_file.parents[2]
    candidates = [
        service_file.parent / "assets" / "j2w-logo.png",
        repo_root / "j2w-flow-insight" / "src" / "assets" / "j2w-logo.png",
    ]
    for candidate in candidates:
        if candidate.exists() and candidate.is_file():
            return candidate
    return None


def _write_docx_report(path: Path, reports: list[dict[str, Any]]) -> None:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    logo_path = _resolve_j2w_logo_path()

    def add_header() -> None:
        if logo_path:
            doc.add_picture(str(logo_path), width=Inches(0.9))
        doc.add_paragraph("J2W Delivery Tracker")

    for idx, report in enumerate(reports):
        if idx > 0:
            doc.add_page_break()
        add_header()
        doc.add_heading(f"Project Status Report: {report.get('project_name') or 'Unnamed Project'}", level=1)
        doc.add_paragraph(
            f"Reporting Window: {report.get('window_start', '')} to {report.get('window_end', '')}"
        )
        doc.add_heading("Executive Summary", level=2)
        doc.add_paragraph(str(report.get("executive_summary") or ""))

        doc.add_heading("Milestone Timeline Stories", level=2)
        for milestone in report.get("milestones") or []:
            title = (
                f"{milestone.get('name', 'Milestone')}  "
                f"[Status: {milestone.get('status') or 'Unknown'}]"
            )
            doc.add_paragraph(title, style="List Bullet")
            doc.add_paragraph(str(milestone.get("timeline_summary") or "No in-window activity."))

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def _write_pdf_report(path: Path, reports: list[dict[str, Any]]) -> None:
    import fitz

    page_width = 595
    page_height = 842
    left_margin = 40
    right_margin = 555
    body_top = 100
    body_bottom = 790
    line_height = 15

    doc = fitz.open()
    page = doc.new_page(width=page_width, height=page_height)
    x = left_margin
    y = body_top
    logo_path = _resolve_j2w_logo_path()
    logo_bytes = logo_path.read_bytes() if logo_path else None

    def _safe_text(target_page: fitz.Page, pos: tuple, text: str, size: int, font: str, color: tuple):
        try:
            target_page.insert_text(pos, text, fontsize=size, fontname=font, color=color)
        except Exception:
            try:
                target_page.insert_text(pos, text, fontsize=size, fontname="helv", color=color)
            except Exception:
                pass

    def _safe_textbox(
        target_page: fitz.Page,
        rect: fitz.Rect,
        text: str,
        size: int,
        font: str,
        color: tuple,
        align: int = 0,
    ) -> None:
        try:
            target_page.insert_textbox(rect, text, fontsize=size, fontname=font, color=color, align=align)
        except Exception:
            try:
                target_page.insert_textbox(rect, text, fontsize=size, fontname="helv", color=color, align=align)
            except Exception:
                pass

    def draw_page_watermark(target_page: fitz.Page) -> None:
        if not logo_bytes:
            return
        try:
            rect = fitz.Rect(130, 220, 465, 600)
            # Preserve original PNG transparency and reduce intensity with a white veil.
            target_page.insert_image(rect, stream=logo_bytes, keep_proportion=True, overlay=False)
            target_page.draw_rect(rect, color=None, fill=(1, 1, 1), fill_opacity=0.82, overlay=True)
        except Exception:
            pass

    def draw_page_header(target_page: fitz.Page, proj_name: str) -> None:
        target_page.draw_rect(fitz.Rect(0, 0, page_width, 60), color=(0.97, 0.98, 1.0), fill=(0.97, 0.98, 1.0), overlay=False)
        header_y = 35
        if logo_bytes:
            try:
                target_page.insert_image(fitz.Rect(x, 15, x + 35, 45), stream=logo_bytes, keep_proportion=True)
                _safe_text(target_page, (x + 42, header_y), "J2W Delivery Tracker", 14, "hebo", (0.1, 0.2, 0.45))
            except Exception:
                _safe_text(target_page, (x, header_y), "J2W Delivery Tracker", 14, "hebo", (0.1, 0.2, 0.45))
        else:
            _safe_text(target_page, (x, header_y), "J2W Delivery Tracker", 14, "hebo", (0.1, 0.2, 0.45))

        _safe_textbox(
            target_page,
            fitz.Rect(360, 24, right_margin, 42),
            f"Generated: {datetime.now().strftime('%b %d, %Y')}",
            8,
            "helv",
            (0.5, 0.5, 0.5),
            align=2,
        )
        target_page.draw_line(fitz.Point(x, 55), fitz.Point(right_margin, 55), color=(0.85, 0.88, 0.92), width=1.0)
        draw_page_watermark(target_page)

    def ensure_space(required_height: int) -> None:
        nonlocal page, y
        if y + required_height <= body_bottom:
            return
        page = doc.new_page(width=page_width, height=page_height)
        draw_page_header(page, "Continued")
        y = body_top

    def write_line(text: str, font_size: int = 10, font_name: str = "helv", color: tuple = (0.2, 0.2, 0.2), indent: int = 0) -> None:
        nonlocal page, y
        if y > body_bottom:
            page = doc.new_page(width=page_width, height=page_height)
            draw_page_header(page, "Continued")
            y = body_top
        _safe_text(page, (x + indent, y), text, font_size, font_name, color)
        y += line_height

    def write_wrapped(
        text: str,
        width: int = 88,
        font_size: float = 10.0,
        color: tuple = (0.2, 0.2, 0.2),
        indent: int = 0,
        prefix: str = "",
        hanging_indent: int = 12,
    ) -> None:
        content = (text or "").strip()
        if not content:
            return

        words = content.split()
        current = ""
        first_line = True

        for w in words:
            candidate = f"{current}{w} "
            if len(candidate) <= width:
                current = candidate
            else:
                if first_line and prefix:
                    write_line(f"{prefix}{current.rstrip()}", font_size=font_size, color=color, indent=indent)
                else:
                    cont_indent = indent + (hanging_indent if prefix else 0)
                    write_line(current.rstrip(), font_size=font_size, color=color, indent=cont_indent)
                first_line = False
                current = w + " "

        if current.strip():
            if first_line and prefix:
                write_line(f"{prefix}{current.rstrip()}", font_size=font_size, color=color, indent=indent)
            else:
                cont_indent = indent + (hanging_indent if prefix else 0)
                write_line(current.rstrip(), font_size=font_size, color=color, indent=cont_indent)

    def normalize_timeline_points(story: str) -> list[str]:
        points: list[str] = []
        for raw in (story or "").splitlines():
            line = (raw or "").strip()
            if not line:
                continue
            if line.startswith("- "):
                line = line[2:].strip()

            lower = line.lower()
            if lower.startswith("schedule:"):
                rest = line.split(":", 1)[1].strip()
                points.append("Schedule")
                for seg in re.split(r",\s+", rest.replace(";", ",")):
                    seg = seg.strip()
                    if seg:
                        points.append(f"  {seg}")
                continue

            points.append(line)
        return points

    for ridx, report in enumerate(reports):
        if ridx > 0:
            page = doc.new_page(width=page_width, height=page_height)
            y = body_top
        else:
            y = body_top

        draw_page_header(page, report.get("project_name", ""))

        # Project Title Section
        y = 110
        ensure_space(110)
        _safe_text(page, (x, y), "PROJECT STATUS REPORT", 9, "hebo", (0.4, 0.4, 0.4))
        y += 24
        _safe_text(page, (x, y), str(report.get("project_name", "Unnamed Project")).upper(), 22, "hebo", (0.0, 0.2, 0.6))
        y += 18
        _safe_text(page, (x, y), f"Reporting Horizon: {report.get('window_start')} to {report.get('window_end')}", 10, "helv", (0.5, 0.5, 0.5))

        y += 44

        # Executive Summary
        ensure_space(80)
        _safe_text(page, (x, y), "EXECUTIVE NARRATIVE", 12, "hebo", (0.1, 0.5, 0.3))
        y += 18
        summary = str(report.get("executive_summary") or "No narrative generated.")
        for chunk in summary.splitlines():
            clean_chunk = (chunk or "").strip()
            if not clean_chunk:
                continue

            if clean_chunk.startswith("##"):
                section_title = clean_chunk.lstrip("#").strip()
                ensure_space(24)
                write_line(section_title, font_size=11, font_name="hebo", color=(0.12, 0.25, 0.55))
                y += 4
                continue

            if clean_chunk.startswith("- "):
                ensure_space(24)
                write_wrapped(
                    clean_chunk[2:].strip(),
                    width=86,
                    font_size=10.5,
                    color=(0.15, 0.15, 0.15),
                    indent=0,
                    prefix="- ",
                    hanging_indent=10,
                )
            else:
                ensure_space(24)
                write_wrapped(clean_chunk, width=90, font_size=10.5, color=(0.15, 0.15, 0.15), indent=0)

            y += 4

        y += 28

        # Milestones
        ensure_space(70)
        _safe_text(page, (x, y), "MILESTONE PROGRESS STORIES", 12, "hebo", (0.3, 0.3, 0.3))
        y += 20

        for m in report.get("milestones", []):
            m_name = str(m.get("name") or "M")
            m_status = str(m.get("status") or "Unknown")

            status_color = (0.2, 0.6, 0.2) if "complete" in m_status.lower() else (0.8, 0.5, 0.0)
            if "risk" in m_status.lower() or "block" in m_status.lower():
                status_color = (0.8, 0.1, 0.1)

            ensure_space(48)
            write_wrapped(
                f"{m_name}: {m.get('description') or ''}",
                width=86,
                font_size=10.5,
                color=(0.1, 0.1, 0.1),
                indent=0,
            )
            y += 2
            _safe_text(page, (x + 15, y), f"STATUS: {m_status.upper()}", 8, "hebo", status_color)
            y += 15

            story = str(m.get("timeline_summary") or "No in-window activity updates recorded.")
            for item in normalize_timeline_points(story):
                if not item:
                    continue

                if item == "Schedule":
                    ensure_space(24)
                    write_line("- Schedule", font_size=10, font_name="hebo", color=(0.30, 0.30, 0.30), indent=15)
                    y += 2
                    continue

                if item.startswith("  "):
                    ensure_space(24)
                    write_wrapped(
                        item.strip(),
                        width=78,
                        font_size=9.5,
                        color=(0.35, 0.35, 0.35),
                        indent=30,
                        prefix="- ",
                        hanging_indent=10,
                    )
                else:
                    ensure_space(24)
                    write_wrapped(
                        item,
                        width=84,
                        font_size=10,
                        color=(0.35, 0.35, 0.35),
                        indent=15,
                        prefix="- ",
                        hanging_indent=10,
                    )
                y += 3

            y += 18

    total_pages = len(doc)
    for pidx in range(total_pages):
        footer_page = doc[pidx]
        footer_page.draw_line(fitz.Point(left_margin, 806), fitz.Point(right_margin, 806), color=(0.9, 0.9, 0.92), width=0.8)
        _safe_text(footer_page, (left_margin, 822), "Confidential - Internal Delivery Management", 7, "helv", (0.55, 0.55, 0.55))
        _safe_text(footer_page, (right_margin - 32, 822), f"{pidx + 1}/{total_pages}", 7, "helv", (0.55, 0.55, 0.55))

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    doc.close()


def export_status_report(
    payload: dict[str, Any],
) -> dict[str, Any]:
    reports = payload.get("reports")
    if not isinstance(reports, list) or not reports:
        raise HTTPException(status_code=400, detail="reports must be a non-empty list")

    batch_mode = str(payload.get("batch_mode") or "single").strip().lower()
    if batch_mode not in {"single", "per_project"}:
        raise HTTPException(status_code=400, detail="batch_mode must be 'single' or 'per_project'")

    export_format = str(payload.get("format") or "docx").strip().lower()
    if export_format not in {"docx", "pdf"}:
        raise HTTPException(status_code=400, detail="format must be 'docx' or 'pdf'")

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    output_dir = UPLOADS_DIR / "reports"
    output_dir.mkdir(parents=True, exist_ok=True)

    generated_files: list[dict[str, Any]] = []

    def _render_file(target_path: Path, report_list: list[dict[str, Any]]) -> None:
        if export_format == "docx":
            _write_docx_report(target_path, report_list)
        else:
            _write_pdf_report(target_path, report_list)

    if batch_mode == "single":
        base = _safe_filename_part(payload.get("file_name") or f"status_report_{timestamp}")
        file_name = f"{base}.{export_format}"
        path = output_dir / file_name
        _render_file(path, reports)
        generated_files.append(
            {
                "project_id": None,
                "project_name": "Unified",
                "file_name": file_name,
                "path": f"uploads/reports/{file_name}",
                "download_url": f"/api/files/uploads/reports/{file_name}?download=true",
            }
        )
    else:
        for report in reports:
            project_name = _safe_filename_part(str(report.get("project_name") or "project"))
            project_id = _safe_filename_part(str(report.get("project_id") or "unknown"))
            file_name = f"{project_name}_{project_id}_{timestamp}.{export_format}"
            path = output_dir / file_name
            _render_file(path, [report])
            generated_files.append(
                {
                    "project_id": report.get("project_id"),
                    "project_name": report.get("project_name"),
                    "file_name": file_name,
                    "path": f"uploads/reports/{file_name}",
                    "download_url": f"/api/files/uploads/reports/{file_name}?download=true",
                }
            )

    return {
        "batch_mode": batch_mode,
        "format": export_format,
        "generated_at": utc_now_iso(),
        "files": generated_files,
    }
