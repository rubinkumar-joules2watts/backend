from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
import logging
from logging.handlers import RotatingFileHandler
from pathlib import Path
import re
from typing import Any
import zipfile

import fitz
import pdfplumber
from bs4 import BeautifulSoup
from docx import Document
from fastapi import HTTPException


MAX_FILE_SIZE_BYTES = 15 * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".html", ".htm"}


def _get_parser_logger() -> logging.Logger:
    logger = logging.getLogger("document_parser")
    if logger.handlers:
        return logger

    logs_dir = Path(__file__).resolve().parent.parent / "logs"
    logs_dir.mkdir(parents=True, exist_ok=True)

    logger.setLevel(logging.INFO)
    formatter = logging.Formatter(
        "%(asctime)s | %(levelname)s | %(name)s | %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )

    file_handler = RotatingFileHandler(
        logs_dir / "document_parser.log",
        maxBytes=2 * 1024 * 1024,
        backupCount=5,
        encoding="utf-8",
    )
    file_handler.setFormatter(formatter)

    stream_handler = logging.StreamHandler()
    stream_handler.setFormatter(formatter)

    logger.addHandler(file_handler)
    logger.addHandler(stream_handler)
    logger.propagate = False
    return logger


logger = _get_parser_logger()

# Reduce noisy third-party PDF warnings in terminal logs.
logging.getLogger("pdfminer").setLevel(logging.ERROR)


@dataclass
class TextBlock:
    text: str
    page: int | None = None
    kind: str = "paragraph"
    x0: float = 0.0  # left edge of the PDF block — used for two-column detection


def _detect_file_type(filename: str, file_bytes: bytes) -> str:
    # Prefer magic-byte sniffing over extension for robustness (e.g., mislabeled .docx files).
    if file_bytes.startswith(b"%PDF-"):
        return "pdf"
    if file_bytes.startswith(b"PK"):
        try:
            with zipfile.ZipFile(BytesIO(file_bytes)) as zf:
                names = set(zf.namelist())
                if "[Content_Types].xml" in names and any(name.startswith("word/") for name in names):
                    return "docx"
        except zipfile.BadZipFile:
            pass

    lowered = filename.lower()
    if lowered.endswith(".pdf"):
        return "pdf"
    if lowered.endswith(".docx"):
        return "docx"
    if lowered.endswith(".html") or lowered.endswith(".htm"):
        return "html"
    return "unsupported"


def _is_heading(text: str) -> bool:
    clean = text.strip()
    if not clean:
        return False
    if len(clean) <= 70 and clean.isupper() and len(clean.split()) <= 10:
        return True
    if re.match(r"^(\d+(\.\d+)*\s+)?[A-Z][\w\s\-/&]{2,80}$", clean) and len(clean.split()) <= 8:
        return True
    if clean.endswith(":") and len(clean.split()) <= 8:
        return True
    return False


def _build_sections(blocks: list[TextBlock]) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    current_title = "Overview"
    current_content: list[str] = []

    for block in blocks:
        txt = block.text.strip()
        if not txt:
            continue
        if _is_heading(txt):
            if current_content:
                sections.append({"title": current_title, "content": current_content})
            current_title = txt.rstrip(":")
            current_content = []
            continue
        current_content.append(txt)

    if current_content:
        sections.append({"title": current_title, "content": current_content})

    return sections


def _extract_html(file_bytes: bytes) -> tuple[list[TextBlock], list[list[list[str]]], dict[str, Any]]:
    soup = BeautifulSoup(file_bytes, "html.parser")

    blocks: list[TextBlock] = []
    for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li"]):
        text = tag.get_text(" ", strip=True)
        if not text:
            continue
        kind = "heading" if tag.name and tag.name.startswith("h") else "paragraph"
        blocks.append(TextBlock(text=text, page=1, kind=kind))

    tables: list[list[list[str]]] = []
    for table in soup.find_all("table"):
        rows: list[list[str]] = []
        for tr in table.find_all("tr"):
            cols = tr.find_all(["th", "td"])
            row = [c.get_text(" ", strip=True) for c in cols]
            if any(row):
                rows.append(row)
        if rows:
            tables.append(rows)

    return blocks, tables, {"pages": 1}


def _extract_docx(file_bytes: bytes) -> tuple[list[TextBlock], list[list[list[str]]], dict[str, Any]]:
    doc = Document(BytesIO(file_bytes))

    blocks: list[TextBlock] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style_name = (para.style.name or "").lower() if para.style else ""
        kind = "heading" if "heading" in style_name else "paragraph"
        blocks.append(TextBlock(text=text, page=1, kind=kind))

    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            vals = [cell.text.strip() for cell in row.cells]
            if any(vals):
                rows.append(vals)
        if rows:
            tables.append(rows)

    return blocks, tables, {"pages": 1}


def _extract_pdf(file_bytes: bytes) -> tuple[list[TextBlock], list[list[list[str]]], dict[str, Any]]:
    blocks: list[TextBlock] = []
    tables: list[list[list[str]]] = []

    with fitz.open(stream=file_bytes, filetype="pdf") as pdf_doc:
        for page_idx, page in enumerate(pdf_doc, start=1):
            raw_blocks = page.get_text("blocks")
            sorted_blocks = sorted(raw_blocks, key=lambda b: (round(b[1], 1), round(b[0], 1)))
            for raw in sorted_blocks:
                text = (raw[4] or "").strip()
                if not text:
                    continue
                block_x0: float = float(raw[0])  # left edge of this text block
                for line in [segment.strip() for segment in text.splitlines() if segment.strip()]:
                    blocks.append(TextBlock(text=line, page=page_idx, kind="paragraph", x0=block_x0))

    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables() or []:
                rows = [[(cell or "").strip() for cell in row] for row in table if row]
                if rows:
                    tables.append(rows)

    page_count = max((b.page or 0) for b in blocks) if blocks else 0
    return blocks, tables, {"pages": page_count}


def _clean_line(text: str) -> str:
    text = text.replace("\x00", " ")
    return re.sub(r"\s+", " ", text).strip(" -\t")


def _find_key_value(lines: list[str], key: str) -> str | None:
    pattern = re.compile(rf"^{re.escape(key)}\s*[:\-]\s*(.+)$", re.IGNORECASE)
    for line in lines:
        match = pattern.match(line.strip())
        if match:
            return match.group(1).strip()
    return None


def _extract_project_metadata(lines: list[str], sections: list[dict[str, Any]]) -> tuple[str, str | None]:
    project_name = _find_key_value(lines, "Project Name")
    client = _find_key_value(lines, "Client")

    if not project_name:
        project_like = [s["title"] for s in sections if s["title"].lower() not in {"overview", "scope"}]
        project_name = project_like[0] if project_like else (lines[0] if lines else "New Project")

    if not client:
        for line in lines[:60]:
            if "client" in line.lower() and ":" in line:
                client = line.split(":", 1)[1].strip() or None
                break

    return project_name.strip(), client.strip() if isinstance(client, str) else None


def _find_date(value: str) -> str | None:
    candidates = [
        r"\b\d{4}-\d{2}-\d{2}\b",
        r"\b\d{1,2}/\d{1,2}/\d{4}\b",
        r"\b\d{1,2}\s+[A-Za-z]{3,9}\s+\d{4}\b",
    ]
    for pattern in candidates:
        match = re.search(pattern, value)
        if match:
            raw = match.group(0)
            for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y", "%d %B %Y", "%d %b %Y"):
                try:
                    return datetime.strptime(raw, fmt).strftime("%Y-%m-%d")
                except ValueError:
                    continue
    return None


_MILESTONE_SECTION_KEYWORDS = ["milestone", "timeline", "delivery plan", "project plan", "delivery schedule"]
_MILESTONE_CONTENT_KEYWORDS = ["milestone", "phase", "deliverable", "go-live", "launch", "release", "handover"]


def _extract_milestones(lines: list[str], sections: list[dict[str, Any]]) -> list[dict[str, Any]]:
    scope_lines: list[str] = []
    for section in sections:
        title = section.get("title", "").lower()
        if any(word in title for word in _MILESTONE_SECTION_KEYWORDS):
            scope_lines.extend(section.get("content", []))

    if not scope_lines:
        scope_lines = [
            line for line in lines
            if any(k in line.lower() for k in _MILESTONE_CONTENT_KEYWORDS)
        ]

    milestones: list[dict[str, Any]] = []
    for idx, line in enumerate(scope_lines, start=1):
        clean = _clean_line(line)
        if len(clean) < 5:
            continue

        start_date = _find_date(clean)
        end_date = None
        if " to " in clean.lower() or " - " in clean:
            parts = re.split(r"\s+to\s+|\s+-\s+", clean)
            if len(parts) > 1:
                end_date = _find_date(parts[-1])

        # Skip lines with no date and no milestone-related keyword — avoids picking up
        # stray headings like "Policy No." from a "Policy Schedule" section.
        has_milestone_keyword = any(k in clean.lower() for k in _MILESTONE_CONTENT_KEYWORDS)
        if not start_date and not has_milestone_keyword:
            continue

        name = re.sub(r"\b\d{4}-\d{2}-\d{2}\b", "", clean)
        name = re.sub(r"\b\d{1,2}/\d{1,2}/\d{4}\b", "", name)
        name = re.sub(r"\b\d{1,2}\s+[A-Za-z]{3,9}\s+\d{4}\b", "", name)
        name = _clean_line(name)
        if not name:
            name = f"Milestone {idx}"

        milestones.append(
            {
                "id": f"m{idx}",
                "name": name[:120],
                "startDate": start_date,
                "endDate": end_date or start_date,
            }
        )

    deduped: list[dict[str, Any]] = []
    seen_names: set[str] = set()
    for item in milestones:
        key = item["name"].lower()
        if key in seen_names:
            continue
        seen_names.add(key)
        deduped.append(item)

    return deduped[:12]


def _parse_month_day(value: str, year: int | None = None) -> str | None:
    value = _clean_line(value)
    if not value:
        return None

    if year is None:
        year_match = re.search(r"\b(20\d{2})\b", value)
        year = int(year_match.group(1)) if year_match else datetime.utcnow().year

    for fmt in ("%d %b %Y", "%d %B %Y"):
        try:
            return datetime.strptime(f"{value} {year}", fmt).strftime("%Y-%m-%d")
        except ValueError:
            continue
    return _find_date(value)


def _extract_timeline_milestones(lines: list[str]) -> list[dict[str, Any]]:
    start_idx = -1
    for i, line in enumerate(lines):
        if "project timeline & milestones" in line.lower():
            start_idx = i
            break
    if start_idx < 0:
        return []

    timeline_lines: list[str] = []
    for line in lines[start_idx : start_idx + 260]:
        lower = line.lower()
        if "success metrics" in lower or "risks, dependencies" in lower:
            break
        if "explore our developer-friendly html to pdf api" in lower:
            continue
        timeline_lines.append(line)

    milestones: list[dict[str, Any]] = []
    i = 0
    while i < len(timeline_lines):
        line = timeline_lines[i]
        if not re.match(r"^Wk\s*\d+", line, flags=re.IGNORECASE):
            i += 1
            continue

        block: list[str] = [line]
        i += 1
        while i < len(timeline_lines) and not re.match(r"^Wk\s*\d+", timeline_lines[i], flags=re.IGNORECASE):
            block.append(timeline_lines[i])
            i += 1

        code = ""
        name_parts: list[str] = []
        date_parts: list[str] = []
        description_parts: list[str] = []
        in_name = False

        for part in block:
            milestone_match = re.match(r"^(M\d+)\s*[:\-]?\s*(.*)$", part, flags=re.IGNORECASE)
            if milestone_match:
                code = milestone_match.group(1).upper()
                maybe_name = milestone_match.group(2).strip()
                if maybe_name:
                    name_parts.append(maybe_name)
                in_name = True
                continue

            if in_name and not date_parts and not re.search(r"\b\d{1,2}\s+[A-Za-z]{3,9}\b", part):
                name_parts.append(part)
                continue

            if re.search(r"\b\d{1,2}\s+[A-Za-z]{3,9}\b", part):
                date_parts.append(part)
                in_name = False
                continue

            if part.lower() in {"period", "milestone", "dates", "description"}:
                continue

            if code:
                description_parts.append(part)

        if not code:
            continue

        name = _clean_line(" ".join(name_parts)) or code
        date_text = " ".join(date_parts)
        range_match = re.search(r"(\d{1,2}\s+[A-Za-z]{3,9})\s*[–-]\s*(\d{1,2}\s+[A-Za-z]{3,9})(?:\s+(20\d{2}))?", date_text)
        if range_match:
            yr = int(range_match.group(3)) if range_match.group(3) else None
            start_date = _parse_month_day(range_match.group(1), yr)
            end_date = _parse_month_day(range_match.group(2), yr)
        else:
            start_date = _find_date(date_text)
            end_date = start_date

        if not start_date:
            continue

        description = _clean_line(" ".join(description_parts))
        if description and description.lower().startswith("all ") and name.lower() in description.lower():
            description = description

        milestones.append(
            {
                "id": code.lower(),
                "name": name[:120],
                "startDate": start_date,
                "endDate": end_date or start_date,
                "description": description,
            }
        )

    deduped: list[dict[str, Any]] = []
    seen: set[str] = set()
    for item in milestones:
        key = item["id"]
        if key in seen:
            continue
        seen.add(key)
        deduped.append(item)
    return deduped[:12]


NOISE_MARKERS = [
    "explore our developer-friendly html to pdf api",
    "printed using pdfcrowd",
    "html to pdf",
]

STOP_ROLE_TITLES = {
    "overview",
    "general notes",
    "description",
    "scope of work",
    "project timeline & milestones",
    "team structure & people deployed",
    "team engagement model & commitments",
    "team engagement model",
    "work location",
    "leave & availability",
    "continuity guarantee",
    "response commitment",
    "customer acknowledgement & sign-off",
}

ABBREVIATION_TOKENS = {"pm", "ar", "sp", "sd", "do", "ba", "qa"}


def _contains_noise(text: str) -> bool:
    lowered = text.lower()
    return any(marker in lowered for marker in NOISE_MARKERS)


def _normalize_role_text(text: str) -> str:
    role = _clean_line(text)
    parts = role.split(maxsplit=1)
    if len(parts) == 2 and parts[0].lower() in ABBREVIATION_TOKENS:
        role = parts[1]
    role = re.sub(r"\s+", " ", role)
    return role.strip(" :-")


_ROLE_SECTION_KEYWORDS = ["resource", "team", "staff", "role", "people"]

_KNOWN_ROLE_TOKENS = {
    "manager", "engineer", "developer", "architect", "analyst", "consultant",
    "lead", "specialist", "administrator", "designer", "coordinator", "director",
    "scrum", "devops", "qa", "tester", "pm",
}


def _is_role_line(text: str) -> bool:
    clean = _normalize_role_text(text)
    if not clean:
        return False

    # Reject sentence fragments: role titles don't end with periods or commas
    if clean.endswith(".") or clean.endswith(","):
        return False

    lowered = clean.lower()
    if lowered in STOP_ROLE_TITLES:
        return False
    if lowered in ABBREVIATION_TOKENS:
        return False
    if _contains_noise(clean):
        return False
    if re.search(r"\d", clean):
        return False

    words = clean.split()
    if len(words) == 0 or len(words) > 6:
        return False

    # Must contain at least one known role-related token to qualify
    has_role_token = any(token in lowered for token in _KNOWN_ROLE_TOKENS)

    if ":" in clean and len(words) <= 6:
        return has_role_token

    if clean.isupper() and len(words) <= 6:
        return has_role_token

    title_like = sum(1 for word in words if word[:1].isupper())
    if title_like >= max(1, len(words) - 1):
        return has_role_token

    return False


def _trim_noise_from_responsibility(text: str) -> str:
    cleaned = _clean_line(text)
    lower = cleaned.lower()
    for marker in ["oversight:", "team engagement model", "work location"] + NOISE_MARKERS:
        idx = lower.find(marker)
        if idx > 0:
            cleaned = _clean_line(cleaned[:idx])
            lower = cleaned.lower()
    return cleaned


def _infer_skills(role: str, responsibility_text: str) -> list[str]:
    text = f"{role} {responsibility_text}".lower()
    keyword_map = {
        "python": "Python",
        "devops": "DevOps",
        "ci/cd": "CI/CD",
        "jenkins": "Jenkins",
        "artifactory": "Artifactory",
        "sbom": "SBOM",
        "pipeline": "Pipeline Automation",
        "monitoring": "Monitoring",
        "testing": "Testing",
        "qa": "QA",
        "architecture": "Architecture",
        "analysis": "Business Analysis",
        "requirements": "Requirements",
        "governance": "Governance",
    }

    skills: list[str] = []
    for key, label in keyword_map.items():
        if key in text and label not in skills:
            skills.append(label)

    role_tokens = [t for t in re.findall(r"[A-Za-z][A-Za-z+/#-]{1,}", role) if len(t) > 2]
    for token in role_tokens:
        normalized = token.upper() if token.upper() in {"QA", "CI/CD"} else token.title()
        if normalized not in skills:
            skills.append(normalized)

    return skills[:8]


def _dedupe_resources_by_role(resources: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Deduplicate resources by canonical role, keeping the richest extracted entry."""
    best_by_role: dict[str, dict[str, Any]] = {}

    for resource in resources:
        role = _clean_line(str(resource.get("role", "")))
        if not role:
            continue

        key = role.lower()
        current = {
            **resource,
            "role": role,
            "skills": [s for s in resource.get("skills", []) if s],
            "responsibilities": _clean_line(str(resource.get("responsibilities", ""))),
            "bandwidth": int(resource.get("bandwidth", 100) or 100),
        }

        existing = best_by_role.get(key)
        if not existing:
            best_by_role[key] = current
            continue

        # Prefer richer extraction: longer responsibility text, then more skills.
        existing_score = (len(existing.get("responsibilities", "")), len(existing.get("skills", [])))
        current_score = (len(current.get("responsibilities", "")), len(current.get("skills", [])))

        if current_score > existing_score:
            merged_skills = existing.get("skills", []) + [
                s for s in current.get("skills", []) if s not in existing.get("skills", [])
            ]
            current["skills"] = merged_skills[:8]
            best_by_role[key] = current
        else:
            existing_skills = existing.get("skills", []) + [
                s for s in current.get("skills", []) if s not in existing.get("skills", [])
            ]
            existing["skills"] = existing_skills[:8]

    deduped: list[dict[str, Any]] = []
    for idx, item in enumerate(best_by_role.values(), start=1):
        item["id"] = f"r{idx}"
        deduped.append(item)
    return deduped


def _parse_team_column(lines: list[str]) -> list[dict[str, Any]]:
    """Parse role/responsibility pairs from a single column of text lines."""
    resources: list[dict[str, Any]] = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not _is_role_line(line):
            i += 1
            continue

        role = _normalize_role_text(line)
        if not role:
            i += 1
            continue

        i += 1
        responsibility_parts: list[str] = []
        while i < len(lines):
            candidate = lines[i].strip()
            if _is_role_line(candidate):
                break
            if candidate.lower() in ABBREVIATION_TOKENS or _contains_noise(candidate):
                i += 1
                continue
            responsibility_parts.append(candidate)
            i += 1

        responsibility = _trim_noise_from_responsibility(" ".join(responsibility_parts))
        resources.append(
            {
                "id": f"r{len(resources)+1}",
                "role": role[:80],
                "skills": _infer_skills(role, responsibility),
                "responsibilities": responsibility[:220],
                "bandwidth": 100,
            }
        )

    return resources


def _extract_team_resources(blocks: list[TextBlock]) -> list[dict[str, Any]]:
    """Extract team resources from 'Team Structure & People Deployed' section.

    Handles two-column PDF layouts by inspecting block x-coordinates and
    processing each column independently before merging.
    """
    # Find the start of the section
    start_idx = -1
    for i, block in enumerate(blocks):
        if "team structure & people deployed" in block.text.lower():
            start_idx = i
            break
    if start_idx < 0:
        return []

    # Collect section blocks until team engagement model or page number
    team_blocks: list[TextBlock] = []
    for block in blocks[start_idx : start_idx + 300]:
        lower = block.text.lower()
        if "team engagement model" in lower or re.match(r"^\d{2}$", block.text.strip()):
            break
        if "explore our developer-friendly html to pdf api" in lower:
            continue
        team_blocks.append(block)

    if not team_blocks:
        return []

    # ── Column detection ────────────────────────────────────────────────────
    # Blocks from HTML/DOCX have x0=0; skip column detection for those.
    x0_values = [round(b.x0) for b in team_blocks if b.x0 > 1]
    column_split: float | None = None

    if x0_values:
        unique_x = sorted(set(x0_values))
        if len(unique_x) >= 2:
            max_gap, gap_at = 0, 0
            for i in range(len(unique_x) - 1):
                gap = unique_x[i + 1] - unique_x[i]
                if gap > max_gap:
                    max_gap, gap_at = gap, i
            # Only treat as two-column when the gap is substantial (≥80 PDF points)
            if max_gap >= 80:
                column_split = (unique_x[gap_at] + unique_x[gap_at + 1]) / 2

    # ── Extract per column ──────────────────────────────────────────────────
    if column_split is not None:
        left_lines = [b.text for b in team_blocks if b.x0 <= column_split]
        right_lines = [b.text for b in team_blocks if b.x0 > column_split]
        resources = _parse_team_column(left_lines) + _parse_team_column(right_lines)
    else:
        resources = _parse_team_column([b.text for b in team_blocks])

    return _dedupe_resources_by_role(resources)


def _extract_resources(lines: list[str], sections: list[dict[str, Any]]) -> list[dict[str, Any]]:
    scope_lines: list[str] = []
    for section in sections:
        title = section.get("title", "").lower()
        if any(word in title for word in ["resource", "team", "staff", "role"]):
            scope_lines.extend(section.get("content", []))

    if not scope_lines:
        return []

    resources: list[dict[str, Any]] = []
    current_role: str | None = None
    current_parts: list[str] = []

    def _flush_current() -> None:
        nonlocal current_role, current_parts, resources
        if not current_role:
            return
        responsibility = _trim_noise_from_responsibility(" ".join(current_parts))
        resources.append(
            {
                "id": f"r{len(resources)+1}",
                "role": current_role[:80],
                "skills": _infer_skills(current_role, responsibility),
                "responsibilities": responsibility[:220],
                "bandwidth": 100,
            }
        )
        current_role = None
        current_parts = []

    for line in scope_lines:
        clean = _clean_line(line)
        if not clean or _contains_noise(clean):
            continue

        if _is_role_line(clean):
            _flush_current()
            current_role = _normalize_role_text(clean)
            continue

        if current_role:
            current_parts.append(clean)

    _flush_current()

    deduped: list[dict[str, Any]] = []
    seen_roles: set[str] = set()
    for resource in resources:
        key = resource["role"].lower()
        if key in seen_roles:
            continue
        seen_roles.add(key)
        deduped.append(resource)

    return deduped[:10]


def extract_proposal_from_document(file_name: str, file_bytes: bytes) -> dict[str, Any]:
    logger.info("extract_start filename=%s bytes=%s", file_name, len(file_bytes))

    if not file_name:
        logger.warning("extract_reject reason=missing_filename")
        raise HTTPException(status_code=400, detail="File name is required")

    file_type = _detect_file_type(file_name, file_bytes)
    if file_type == "unsupported":
        logger.warning("extract_reject reason=unsupported_type filename=%s", file_name)
        raise HTTPException(status_code=400, detail="Unsupported file type. Use PDF, DOCX, or HTML")

    ext = "." + file_name.lower().rsplit(".", 1)[-1]
    if ext not in ALLOWED_EXTENSIONS:
        logger.warning("extract_reject reason=unsupported_ext ext=%s filename=%s", ext, file_name)
        raise HTTPException(status_code=400, detail="Unsupported file extension")

    if len(file_bytes) == 0:
        logger.warning("extract_reject reason=empty_file filename=%s", file_name)
        raise HTTPException(status_code=400, detail="Uploaded file is empty")

    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        logger.warning("extract_reject reason=oversize filename=%s bytes=%s", file_name, len(file_bytes))
        raise HTTPException(status_code=413, detail="File exceeds 15MB upload limit")

    started_at = datetime.utcnow()

    if file_type == "html":
        blocks, tables, source_meta = _extract_html(file_bytes)
    elif file_type == "docx":
        blocks, tables, source_meta = _extract_docx(file_bytes)
    else:
        blocks, tables, source_meta = _extract_pdf(file_bytes)

    cleaned_lines = [_clean_line(block.text) for block in blocks if _clean_line(block.text)]
    sections = _build_sections(blocks)
    project_name, client = _extract_project_metadata(cleaned_lines, sections)
    timeline_milestones = _extract_timeline_milestones(cleaned_lines)
    generic_milestones = _extract_milestones(cleaned_lines, sections)
    milestones = timeline_milestones or generic_milestones

    team_resources = _extract_team_resources(blocks)
    generic_resources = _extract_resources(cleaned_lines, sections)
    resources = team_resources if team_resources else generic_resources

    warnings: list[str] = []
    if not milestones and not resources:
        warnings.append("INCOMPATIBLE_DOCUMENT: This document does not appear to be a project proposal. No milestones or roles were found. Please add details manually.")
    else:
        if not milestones:
            warnings.append("No clear milestones found. Please review and add them in step 1.")
        if not resources:
            warnings.append("No roles found in document.")

    elapsed_ms = int((datetime.utcnow() - started_at).total_seconds() * 1000)
    logger.info(
        "extract_success filename=%s type=%s sections=%s tables=%s milestones=%s resources=%s warnings=%s ms=%s",
        file_name,
        file_type,
        len(sections),
        len(tables),
        len(milestones),
        len(resources),
        len(warnings),
        elapsed_ms,
    )

    return {
        "metadata": {
            "filename": file_name,
            "file_type": file_type,
            "file_size": len(file_bytes),
            "processing_ms": elapsed_ms,
            "pages": source_meta.get("pages", 0),
        },
        "content": {
            "sections": sections,
            "tables": tables,
            "text_preview": cleaned_lines[:30],
        },
        "proposal": {
            "project_name": project_name or "New Project",
            "client": client,
            "milestones": milestones,
            "resources": resources,
            "warnings": warnings,
        },
    }